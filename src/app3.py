# -*- coding: utf-8 -*-
# Retículo hexagonal 45 m con auxiliares (interior y exterior <= 1 m)
# Exporta CSV, DXF (por capas y rotulado) e informe Word con imágenes y tablas.

import os
import math
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from shapely.geometry import Polygon, Point
from shapely.ops import nearest_points
import numpy.linalg as npl

try:
    import ezdxf
except Exception as e:
    raise RuntimeError("Falta ezdxf. Instala con: pip install ezdxf==1.1.4") from e

try:
    from docx import Document
    from docx.shared import Cm
except Exception as e:
    raise RuntimeError("Falta python-docx. Instala con: pip install python-docx") from e

# -----------------------------
# Parámetros (ajustables)
# -----------------------------
archivo_excel = "poligono.xlsx"            # debe tener columnas X, Y
separacion_m = 45.0                        # separación de la malla hexagonal
incluir_borde = True                       # usar covers (incluye borde) vs contains
min_dist_final = 31.5                      # distancia mínima entre puntos finales
umbral_exterior_valido = 1.0               # auxiliares fuera a <= 1.0 m del borde
margen_interior = 1.0                      # proyección hacia dentro para BordeAux
muestras_offset = 21                       # resolución de búsqueda de offset (x e y)

# Salidas
csv_puntos = "puntos_hex_45m_aux_total.csv"
csv_puntos_2dec = "puntos_finales_tabla_2dec.csv"
csv_vertices = "poligono_vertices.csv"
dxf_capas = "reticulo_poligono_45m_aux_total.dxf"
dxf_rotulado = "reticulo_poligono_45m_final_rotulado.dxf"
fig_interiores = "fig_reticulo_45m_interior.png"
fig_total = "fig_reticulo_45m_aux_total.png"
informe_docx = "informe_reticulo_45m_final_v3.docx"

# -----------------------------
# Utilidades
# -----------------------------
def cargar_poligono(path_excel):
    df = pd.read_excel(path_excel, sheet_name=0)
    df.columns = [str(c).strip() for c in df.columns]
    low = [c.lower() for c in df.columns]
    if "x" not in low or "y" not in low:
        raise ValueError("El Excel debe tener columnas X e Y (sin importar mayúsculas).")
    col_x = df.columns[low.index("x")]
    col_y = df.columns[low.index("y")]
    df = df[[col_x, col_y]].copy()
    df[col_x] = pd.to_numeric(df[col_x], errors="coerce")
    df[col_y] = pd.to_numeric(df[col_y], errors="coerce")
    df = df.dropna()
    vertices = list(df.itertuples(index=False, name=None))
    pol = Polygon(vertices)
    if not pol.is_valid:
        pol = pol.buffer(0)
    return pol, vertices

def generar_pos_hex(centro, d, radio):
    # Retículo hexagonal punto a punto a partir de un centro y radio de cobertura
    n = int(math.ceil(radio / (d * math.sqrt(3) / 2.0))) + 2
    pts = []
    cx, cy = centro
    dy = d * math.sqrt(3) / 2.0
    for i in range(-n, n + 1):
        for j in range(-n, n + 1):
            x = cx + d * (j + (0.5 if (i % 2) != 0 else 0.0))
            y = cy + dy * i
            pts.append((x, y))
    return pts

def optimizar_offset(poligono, d, muestras=21, incluir_borde=True):
    # Busca desplazamiento (ox, oy) que maximice puntos interiores
    centroide = tuple(poligono.centroid.coords)[0]
    coords = np.array(poligono.exterior.coords)[:, :2]
    radio = float(np.max(npl.norm(coords - np.array(centroide), axis=1))) + d
    xs = np.linspace(0.0, d, muestras)
    ys = np.linspace(0.0, d * math.sqrt(3) / 2.0, muestras)
    mejor = -1
    mejor_off = (0.0, 0.0)
    mejores_pts = []
    for oy in ys:
        for ox in xs:
            c = (centroide[0] + ox, centroide[1] + oy)
            pos = generar_pos_hex(c, d, radio)
            if incluir_borde:
                pts = [p for p in pos if poligono.covers(Point(p))]
            else:
                pts = [p for p in pos if poligono.contains(Point(p))]
            if len(pts) > mejor:
                mejor = len(pts)
                mejor_off = (ox, oy)
                mejores_pts = pts
    origen = (centroide[0] + mejor_off[0], centroide[1] + mejor_off[1])
    todas = generar_pos_hex(origen, d, radio)
    return origen, todas

def consolidar_auxiliares(poligono, puntos_interiores, puntos_exteriores, min_dist_final, umbral_exterior_valido, margen_interior):
    # Crea listas: proyectados hacia dentro (BordeAux) y exteriores válidos (BordeAux_Exterior)
    borde = poligono.boundary
    aux_proy = []
    for px, py in puntos_exteriores:
        p = Point(px, py)
        d = p.distance(borde)
        if d <= separacion_m * 0.5:
            q = nearest_points(p, borde)[1]
            vx = poligono.centroid.x - q.x
            vy = poligono.centroid.y - q.y
            norm = math.hypot(vx, vy)
            if norm == 0.0:
                continue
            ux, uy = vx / norm, vy / norm
            cand = (q.x + ux * margen_interior, q.y + uy * margen_interior)
            if poligono.contains(Point(cand)):
                aux_proy.append(cand)

    aux_ext = []
    for px, py in puntos_exteriores:
        p = Point(px, py)
        d = p.distance(borde)
        if d > 0.0 and d <= umbral_exterior_valido:
            aux_ext.append((px, py))

    # Distancia mínima
    def es_valido(cand, existentes, min_d):
        if len(existentes) == 0:
            return True
        arr = np.array(existentes)
        dmin = float(np.min(npl.norm(arr - np.array(cand), axis=1)))
        return dmin >= min_d

    finales = list(puntos_interiores)
    etiquetas = ["Interior"] * len(puntos_interiores)

    for cand in aux_proy:
        if es_valido(cand, finales, min_dist_final):
            finales.append(cand)
            etiquetas.append("BordeAux")

    for cand in aux_ext:
        if es_valido(cand, finales, min_dist_final):
            finales.append(cand)
            etiquetas.append("BordeAux_Exterior")

    return finales, etiquetas

def exportar_dxf(vertices, puntos_finales, etiquetas_finales, nombre_dxf):
    plan = ezdxf.new(setup=True)
    for layer in ["Puntos_Interior", "Puntos_BordeAux", "Puntos_BordeAux_Exterior", "Poligono"]:
        if layer not in plan.layers:
            plan.layers.add(layer)
    ms = plan.modelspace()
    coords = list(vertices)
    if coords[0] != coords[-1]:
        coords.append(coords[0])
    ms.add_lwpolyline(coords, format="xy", dxfattribs={"closed": True, "layer": "Poligono"})
    for (x, y), t in zip(puntos_finales, etiquetas_finales):
        lyr = "Puntos_Interior" if t == "Interior" else ("Puntos_BordeAux" if t == "BordeAux" else "Puntos_BordeAux_Exterior")
        ms.add_point((float(x), float(y)), dxfattribs={"layer": lyr})
    plan.saveas(nombre_dxf)

def exportar_dxf_rotulado(vertices, df_final, nombre_dxf):
    plan = ezdxf.new(setup=True)
    for layer in ["Puntos_Interior", "Puntos_BordeAux", "Puntos_BordeAux_Exterior", "Rotulos_ID", "Poligono"]:
        if layer not in plan.layers:
            plan.layers.add(layer)
    ms = plan.modelspace()
    coords = list(vertices)
    if coords[0] != coords[-1]:
        coords.append(coords[0])
    ms.add_lwpolyline(coords, format="xy", dxfattribs={"closed": True, "layer": "Poligono"})
    for _, r in df_final.iterrows():
        x = float(r["X"])
        y = float(r["Y"])
        t = str(r["Tipo"])
        pid = int(r["ID"])
        lyr = "Puntos_Interior" if t == "Interior" else ("Puntos_BordeAux" if t == "BordeAux" else "Puntos_BordeAux_Exterior")
        ms.add_point((x, y), dxfattribs={"layer": lyr})
        txt = ms.add_text(str(pid), dxfattribs={"height": 1.2, "layer": "Rotulos_ID"})
        txt.set_dxf_attrib("insert", (x + 1.0, y + 1.0))  # desplazado 1 m

    plan.saveas(nombre_dxf)

def graficar(poligono, origen, puntos_interiores, puntos_finales, etiquetas_finales, fig_interiores_path, fig_total_path):
    xp, yp = poligono.exterior.xy

    plt.figure(figsize=(6, 5))
    plt.fill(xp, yp, alpha=0.4, fc="lightblue", ec="darkblue")
    if len(puntos_interiores) > 0:
        pi = np.array(puntos_interiores)
        plt.scatter(pi[:, 0], pi[:, 1], c="green", s=18, label="Interior")
    plt.scatter(origen[0], origen[1], c="red", s=60, label="Punto inicial")
    plt.title("Puntos interiores (45 m)")
    plt.xlabel("X"); plt.ylabel("Y"); plt.axis("equal"); plt.grid(True); plt.legend()
    plt.savefig(fig_interiores_path, dpi=180, bbox_inches="tight")
    plt.close()

    plt.figure(figsize=(6, 5))
    plt.fill(xp, yp, alpha=0.4, fc="lightblue", ec="darkblue")
    if len(puntos_finales) > 0:
        pi = [p for p, t in zip(puntos_finales, etiquetas_finales) if t == "Interior"]
        if len(pi) > 0:
            pi_arr = np.array(pi)
            plt.scatter(pi_arr[:, 0], pi_arr[:, 1], c="green", s=18, label="Interior")
        pa = [p for p, t in zip(puntos_finales, etiquetas_finales) if t == "BordeAux"]
        if len(pa) > 0:
            pa_arr = np.array(pa)
            plt.scatter(pa_arr[:, 0], pa_arr[:, 1], c="orange", s=28, label="BordeAux")
        pe = [p for p, t in zip(puntos_finales, etiquetas_finales) if t == "BordeAux_Exterior"]
        if len(pe) > 0:
            pe_arr = np.array(pe)
            plt.scatter(pe_arr[:, 0], pe_arr[:, 1], c="purple", s=28, marker="x", label="BordeAux_Exterior")
    plt.scatter(origen[0], origen[1], c="red", s=60, label="Punto inicial")
    plt.title("Interiores y auxiliares (45 m)")
    plt.xlabel("X"); plt.ylabel("Y"); plt.axis("equal"); plt.grid(True); plt.legend()
    plt.savefig(fig_total_path, dpi=180, bbox_inches="tight")
    plt.close()

def crear_informe(poligono_vertices, fig_interiores_path, fig_total_path, df_final, informe_path):
    doc = Document()
    doc.add_heading("Informe del retículo hexagonal (45 m)", level=1)

    doc.add_heading("1. Tabla de puntos del polígono", level=2)
    tbl = doc.add_table(rows=1, cols=2)
    hdr = tbl.rows[0].cells
    hdr[0].text = "X"
    hdr[1].text = "Y"
    for x, y in poligono_vertices:
        row = tbl.add_row().cells
        row[0].text = str(x)
        row[1].text = str(y)

    doc.add_heading("2. Distribución de puntos", level=2)
    doc.add_paragraph("Puntos que caen dentro del polígono (Interiores):")
    if os.path.exists(fig_interiores_path):
        doc.add_picture(fig_interiores_path, width=Cm(14))
    doc.add_paragraph("Puntos interiores y auxiliares (BordeAux y BordeAux_Exterior):")
    if os.path.exists(fig_total_path):
        doc.add_picture(fig_total_path, width=Cm(14))

    ni = int((df_final["Tipo"] == "Interior").sum())
    nb = int((df_final["Tipo"] == "BordeAux").sum())
    ne = int((df_final["Tipo"] == "BordeAux_Exterior").sum())
    nt = int(len(df_final))
    doc.add_heading("3. Totales por tipología", level=2)
    doc.add_paragraph("Interior: " + str(ni))
    doc.add_paragraph("BordeAux: " + str(nb))
    doc.add_paragraph("BordeAux_Exterior: " + str(ne))
    doc.add_paragraph("Total: " + str(nt))

    doc.add_heading("4. Tabla de puntos finales (ID, X, Y, Tipo) con 2 decimales", level=2)
    tbl2 = doc.add_table(rows=1, cols=4)
    h2 = tbl2.rows[0].cells
    h2[0].text = "ID"
    h2[1].text = "X"
    h2[2].text = "Y"
    h2[3].text = "Tipo"
    for _, r in df_final.iterrows():
        row = tbl2.add_row().cells
        row[0].text = str(int(r["ID"]))
        row[1].text = format(float(r["X"]), ".2f")
        row[2].text = format(float(r["Y"]), ".2f")
        row[3].text = str(r["Tipo"])

    doc.save(informe_path)

# -----------------------------
# Ejecución
# -----------------------------
poligono, vertices = cargar_poligono(archivo_excel)

# Guardar vértices a CSV
pd.DataFrame(vertices, columns=["X", "Y"]).to_csv(csv_vertices, index=False)

# Optimizar offset y construir malla
origen, pos_todas = optimizar_offset(poligono, separacion_m, muestras=muestras_offset, incluir_borde=incluir_borde)

if incluir_borde:
    puntos_interiores = [p for p in pos_todas if poligono.covers(Point(p))]
    puntos_exteriores = [p for p in pos_todas if not poligono.covers(Point(p))]
else:
    puntos_interiores = [p for p in pos_todas if poligono.contains(Point(p))]
    puntos_exteriores = [p for p in pos_todas if not poligono.contains(Point(p))]

# Auxiliares y consolidación
puntos_finales, etiquetas_finales = consolidar_auxiliares(
    poligono=poligono,
    puntos_interiores=puntos_interiores,
    puntos_exteriores=puntos_exteriores,
    min_dist_final=min_dist_final,
    umbral_exterior_valido=umbral_exterior_valido,
    margen_interior=margen_interior
)

# Orden e IDs: Interior -> BordeAux -> BordeAux_Exterior
orden_tipo = {"Interior": 0, "BordeAux": 1, "BordeAux_Exterior": 2}
filas = []
for (x, y), t in zip(puntos_finales, etiquetas_finales):
    filas.append({"X": float(x), "Y": float(y), "Tipo": t, "orden": orden_tipo.get(t, 99)})
df_final = pd.DataFrame(filas).sort_values(["orden", "X", "Y"]).reset_index(drop=True)
df_final["ID"] = df_final.index + 1

# Exportar CSVs
df_final[["X", "Y", "Tipo"]].to_csv(csv_puntos, index=False)
df_csv2 = df_final.copy()
df_csv2["X"] = df_csv2["X"].map(lambda v: format(float(v), ".2f"))
df_csv2["Y"] = df_csv2["Y"].map(lambda v: format(float(v), ".2f"))
df_csv2[["ID", "X", "Y", "Tipo"]].to_csv(csv_puntos_2dec, index=False)

# Exportar DXF por capas y DXF rotulado
exportar_dxf(vertices, list(df_final[["X", "Y"]].itertuples(index=False, name=None)), list(df_final["Tipo"]), dxf_capas)
exportar_dxf_rotulado(vertices, df_final[["ID", "X", "Y", "Tipo"]], dxf_rotulado)

# Figuras
graficar(poligono, origen, puntos_interiores, list(df_final[["X", "Y"]].itertuples(index=False, name=None)), list(df_final["Tipo"]), fig_interiores, fig_total)

# Informe Word
crear_informe(vertices, fig_interiores, fig_total, df_final[["ID", "X", "Y", "Tipo"]], informe_docx)

print("Listo.")
print("CSV puntos: " + csv_puntos)
print("CSV puntos 2 decimales: " + csv_puntos_2dec)
print("CSV vertices: " + csv_vertices)
print("DXF por capas: " + dxf_capas)
print("DXF rotulado: " + dxf_rotulado)
print("Informe Word: " + informe_docx)
print("Totales -> Interior: " + str((df_final['Tipo'] == 'Interior').sum()) + ", BordeAux: " + str((df_final['Tipo'] == 'BordeAux').sum()) + ", BordeAux_Exterior: " + str((df_final['Tipo'] == 'BordeAux_Exterior').sum()) + ", Total: " + str(len(df_final)))