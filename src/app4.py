# -*- coding: utf-8 -*-
# Script integral para:
# 1) Leer un polígono desde poligono.xlsx (columnas X/Y)
# 2) Generar un retículo hexagonal con distancia fija (p. ej. 45 m) optimizando desplazamiento
# 3) Exportar CSV de puntos, CSV de vértices, DXF y un informe Word con imágenes en centímetros
# 4) Analizar puntos que NO tienen al menos un vecino a la distancia establecida, con segmento dentro del polígono,
#    y que además no están conectados al resto (componente principal).
#
# Requisitos de paquetes:
#   pandas, numpy, shapely, matplotlib, ezdxf, python-docx
#   (opcional para acelerar vecinos): scipy
#
# Nota: No se utilizan f-strings; se concatenan cadenas por compatibilidad.

import os
import math
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from shapely.geometry import Polygon, Point, LineString

# Intenta importar herramientas opcionales
try:
    import ezdxf
    DISPONIBLE_EZDXF = True
except Exception:
    DISPONIBLE_EZDXF = False

try:
    from docx import Document
    from docx.shared import Cm
    DISPONIBLE_DOCX = True
except Exception:
    DISPONIBLE_DOCX = False

try:
    from scipy.spatial import cKDTree as KDTree
    DISPONIBLE_KDTREE = True
except Exception:
    DISPONIBLE_KDTREE = False

# ---------------------------
# Parámetros principales
# ---------------------------
archivo_excel = 'poligono.xlsx'
nombre_hoja = 0
distancia_objetivo_m = 45.0         # Distancia entre puntos del retículo
incluir_borde = True                 # Incluir puntos en el borde del polígono
optimizar_desplazamiento = True      # Buscar el offset que maximiza cantidad de puntos dentro
tolerancia_m = 0.5                   # Tolerancia para comparar distancia de aristas

# Nombres de salida basados en la distancia
nombre_csv_puntos = 'puntos_hex_' + str(int(distancia_objetivo_m)) + 'm.csv'
nombre_csv_vertices = 'poligono_vertices.csv'
nombre_dxf = 'reticulo_poligono_' + str(int(distancia_objetivo_m)) + 'm.dxf'
nombre_informe = 'informe_reticulo_hex_' + str(int(distancia_objetivo_m)) + 'm.docx'
nombre_fig_origen = 'fig_poligono_origen.png'
nombre_fig_puntos = 'fig_poligono_puntos.png'
nombre_csv_problematicos = 'puntos_problematicos_' + str(int(distancia_objetivo_m)) + 'm.csv'
nombre_fig_problematicos = 'puntos_problematicos_' + str(int(distancia_objetivo_m)) + 'm.png'

# ---------------------------
# Funciones auxiliares
# ---------------------------
def cargar_poligono_desde_excel(ruta_excel, hoja):
    df = pd.read_excel(ruta_excel, sheet_name=hoja)
    df.columns = [str(c).strip() for c in df.columns]
    columnas_minus = [c.lower() for c in df.columns]
    if 'x' not in columnas_minus or 'y' not in columnas_minus:
        raise ValueError('El archivo necesita columnas X e Y')
    columna_x = df.columns[columnas_minus.index('x')]
    columna_y = df.columns[columnas_minus.index('y')]
    df_coords = df[[columna_x, columna_y]].copy()
    df_coords[columna_x] = pd.to_numeric(df_coords[columna_x], errors='coerce')
    df_coords[columna_y] = pd.to_numeric(df_coords[columna_y], errors='coerce')
    df_coords = df_coords.dropna()
    vertices = list(df_coords.itertuples(index=False, name=None))
    if len(vertices) < 3:
        raise ValueError('Se requieren al menos 3 vértices válidos')
    poligono = Polygon(vertices)
    if not poligono.is_valid:
        poligono = poligono.buffer(0)
    return df_coords, poligono

def generar_reticulo_hexagonal(poligono, centro, distancia, incluir_borde_local):
    # Calcula un radio de cobertura y genera una malla hexagonal centrada en "centro"
    coords_pol = np.array(poligono.exterior.coords)
    radio = float(np.max(np.linalg.norm(coords_pol[:, :2] - np.array(centro), axis=1)))
    paso_vertical = distancia * math.sqrt(3.0) / 2.0
    filas = int(math.ceil(radio / paso_vertical)) + 2

    puntos_locales = []
    for i in range(-filas, filas + 1):
        for j in range(-filas, filas + 1):
            xi = centro[0] + distancia * (j + (0.5 if (i % 2) != 0 else 0.0))
            yi = centro[1] + paso_vertical * i
            punto = Point(xi, yi)
            adentro = poligono.covers(punto) if incluir_borde_local else poligono.contains(punto)
            if adentro:
                puntos_locales.append((xi, yi))
    return puntos_locales

def optimizar_offset(poligono, distancia, incluir_borde_local):
    centroide = tuple(poligono.centroid.coords)[0]
    paso_vertical = distancia * math.sqrt(3.0) / 2.0
    rangos_x = np.linspace(0.0, distancia, 21)
    rangos_y = np.linspace(0.0, paso_vertical, 21)
    mejor_cantidad = -1
    mejor_despl = (0.0, 0.0)
    mejores_puntos = []
    for dy in rangos_y:
        for dx in rangos_x:
            centro = (centroide[0] + dx, centroide[1] + dy)
            pts = generar_reticulo_hexagonal(poligono, centro, distancia, incluir_borde_local)
            if len(pts) > mejor_cantidad:
                mejor_cantidad = len(pts)
                mejor_despl = (dx, dy)
                mejores_puntos = pts
    origen = (centroide[0] + mejor_despl[0], centroide[1] + mejor_despl[1])
    return origen, mejores_puntos

def graficar_poligono_y_origen(poligono, origen, ruta_salida, titulo):
    xp, yp = poligono.exterior.xy
    plt.figure(figsize=(6, 5))
    plt.fill(xp, yp, alpha=0.4, fc='lightblue', ec='darkblue')
    plt.scatter(origen[0], origen[1], c='red', s=60, label='Punto inicial')
    plt.title(titulo)
    plt.xlabel('X')
    plt.ylabel('Y')
    plt.axis('equal')
    plt.grid(True)
    plt.legend()
    plt.savefig(ruta_salida, dpi=180, bbox_inches='tight')
    plt.show()

def graficar_poligono_y_puntos(poligono, origen, puntos, ruta_salida, titulo):
    xp, yp = poligono.exterior.xy
    plt.figure(figsize=(6, 5))
    plt.fill(xp, yp, alpha=0.4, fc='lightblue', ec='darkblue')
    if len(puntos) > 0:
        arr = np.array(puntos)
        plt.scatter(arr[:, 0], arr[:, 1], c='green', s=15, label='Puntos dentro')
    plt.scatter(origen[0], origen[1], c='red', s=60, label='Punto inicial')
    plt.title(titulo)
    plt.xlabel('X')
    plt.ylabel('Y')
    plt.axis('equal')
    plt.grid(True)
    plt.legend()
    plt.savefig(ruta_salida, dpi=180, bbox_inches='tight')
    plt.show()

def exportar_csv_vertices(vertices, ruta_csv):
    pd.DataFrame(vertices, columns=['X', 'Y']).to_csv(ruta_csv, index=False)

def exportar_csv_puntos(puntos, ruta_csv):
    pd.DataFrame(puntos, columns=['X', 'Y']).to_csv(ruta_csv, index=False)

def exportar_dxf(poligono, puntos, ruta_dxf):
    if not DISPONIBLE_EZDXF:
        print('Aviso: ezdxf no está instalado; omitiendo exportación DXF')
        return
    doc = ezdxf.new(setup=True)
    ms = doc.modelspace()
    coords = list(poligono.exterior.coords)
    if tuple(coords[0]) != tuple(coords[-1]):
        coords.append(coords[0])
    ms.add_lwpolyline(coords, format='xy', dxfattribs={'closed': True})
    for x, y in puntos:
        ms.add_point((x, y))
    doc.saveas(ruta_dxf)

def crear_informe_word(poligono, vertices, origen, puntos, ruta_fig_origen, ruta_fig_puntos, ruta_docx, archivo_fuente, distancia_m):
    if not DISPONIBLE_DOCX:
        print('Aviso: python-docx no está instalado; omitiendo informe Word')
        return
    documento = Document()
    documento.add_heading('Informe de retículo hexagonal en polígono', level=1)

    documento.add_heading('Datos de entrada', level=2)
    documento.add_paragraph('Archivo: ' + archivo_fuente)
    documento.add_paragraph('Vértices del polígono (X, Y):')
    tabla = documento.add_table(rows=1, cols=2)
    tabla.rows[0].cells[0].text = 'X'
    tabla.rows[0].cells[1].text = 'Y'
    for x, y in vertices:
        fila = tabla.add_row().cells
        fila[0].text = str(x)
        fila[1].text = str(y)

    documento.add_paragraph('Punto inicial del retículo (X, Y): ' + str(round(origen[0], 3)) + ', ' + str(round(origen[1], 3)))
    documento.add_paragraph('Polígono y punto inicial:')
    documento.add_picture(ruta_fig_origen, width=Cm(14))

    documento.add_heading('Resultados', level=2)
    documento.add_paragraph('Separación del retículo: ' + str(int(distancia_m)) + ' m')
    documento.add_paragraph('Número de puntos dentro del polígono (incluyendo el inicial): ' + str(len(puntos)))
    documento.add_paragraph('Polígono con los puntos dentro (punto inicial en rojo):')
    documento.add_picture(ruta_fig_puntos, width=Cm(14))
    documento.save(ruta_docx)

def construir_vecindad_valida(poligono, arreglo_puntos, distancia_m, tolerancia):
    # Construye lista de adyacencia donde hay arista entre puntos si:
    # - La distancia está dentro de [distancia - tolerancia, distancia + tolerancia]
    # - El segmento entre los puntos está cubierto por el polígono
    n = arreglo_puntos.shape[0]
    lista_adyacencia = [[] for _ in range(n)]

    if DISPONIBLE_KDTREE:
        arbol = KDTree(arreglo_puntos)
        candidatos = arbol.query_ball_point(arreglo_puntos, distancia_m + tolerancia)
    else:
        candidatos = []
        for i in range(n):
            d = np.linalg.norm(arreglo_puntos - arreglo_puntos[i], axis=1)
            vecinos = list(np.where(d <= distancia_m + tolerancia)[0])
            candidatos.append(vecinos)

    for i in range(n):
        for j in candidatos[i]:
            if j == i:
                continue
            dij = np.linalg.norm(arreglo_puntos[i] - arreglo_puntos[j])
            if dij < distancia_m - tolerancia or dij > distancia_m + tolerancia:
                continue
            segmento = LineString([tuple(arreglo_puntos[i]), tuple(arreglo_puntos[j])])
            if poligono.covers(segmento):
                lista_adyacencia[i].append(j)

    return lista_adyacencia

def componentes_conexas(lista_adyacencia):
    n = len(lista_adyacencia)
    visitado = np.zeros(n, dtype=bool)
    comps = []
    for i in range(n):
        if not visitado[i]:
            cola = [i]
            visitado[i] = True
            comp = [i]
            while len(cola) > 0:
                u = cola.pop(0)
                for v in lista_adyacencia[u]:
                    if not visitado[v]:
                        visitado[v] = True
                        cola.append(v)
                        comp.append(v)
            comps.append(comp)
    return comps

def analizar_puntos_problematicos(poligono, puntos, distancia_m, tolerancia, ruta_csv_salida, ruta_fig_salida):
    arreglo = np.array(puntos)
    if arreglo.shape[0] == 0:
        pd.DataFrame(columns=['X','Y','sin_vecino','fuera_componente_principal']).to_csv(ruta_csv_salida, index=False)
        # Dibujo vacío
        xp, yp = poligono.exterior.xy
        plt.figure(figsize=(6,5))
        plt.fill(xp, yp, alpha=0.35, fc='lightblue', ec='darkblue')
        plt.title('Sin puntos')
        plt.axis('equal')
        plt.grid(True)
        plt.savefig(ruta_fig_salida, dpi=180, bbox_inches='tight')
        plt.show()
        return pd.DataFrame()

    ady = construir_vecindad_valida(poligono, arreglo, distancia_m, tolerancia)
    sin_vecino = [i for i in range(len(ady)) if len(ady[i]) == 0]
    comps = componentes_conexas(ady)
    tamanos = [len(c) for c in comps]
    if len(tamanos) > 0:
        idx_max = int(np.argmax(tamanos))
        comp_principal = set(comps[idx_max])
    else:
        comp_principal = set()

    fuera_comp_principal = [i for i in range(arreglo.shape[0]) if i not in comp_principal]

    mascara_problematicos = np.zeros(arreglo.shape[0], dtype=bool)
    for i in sin_vecino:
        mascara_problematicos[i] = True
    for i in fuera_comp_principal:
        mascara_problematicos[i] = True

    df_problem = pd.DataFrame({
        'X': arreglo[:, 0],
        'Y': arreglo[:, 1],
        'sin_vecino': [i in sin_vecino for i in range(arreglo.shape[0])],
        'fuera_componente_principal': [i in fuera_comp_principal for i in range(arreglo.shape[0])]
    })
    df_problem_only = df_problem[mascara_problematicos].copy()
    df_problem_only.to_csv(ruta_csv_salida, index=False)

    # Gráfico
    xp, yp = poligono.exterior.xy
    plt.figure(figsize=(6,5))
    plt.fill(xp, yp, alpha=0.35, fc='lightblue', ec='darkblue')
    plt.scatter(arreglo[:,0], arreglo[:,1], s=18, c='gray', label='Puntos')
    if df_problem_only.shape[0] > 0:
        plt.scatter(df_problem_only['X'], df_problem_only['Y'], s=45, c='red', label='Problemáticos')
    plt.title('Puntos problemáticos a ' + str(int(distancia_m)) + ' m')
    plt.xlabel('X')
    plt.ylabel('Y')
    plt.axis('equal')
    plt.grid(True)
    plt.legend()
    plt.savefig(ruta_fig_salida, dpi=180, bbox_inches='tight')
    plt.show()

    return df_problem_only

# ---------------------------
# Flujo principal
# ---------------------------
if __name__ == '__main__':
    # 1) Cargar polígono
    df_vertices, poligono = cargar_poligono_desde_excel(archivo_excel, nombre_hoja)
    vertices = list(df_vertices.itertuples(index=False, name=None))

    # 2) Generar retículo hexagonal con offset optimizado (o centroide)
    if optimizar_desplazamiento:
        origen_reticulo, puntos = optimizar_offset(poligono, distancia_objetivo_m, incluir_borde)
    else:
        origen_reticulo = tuple(poligono.centroid.coords)[0]
        puntos = generar_reticulo_hexagonal(poligono, origen_reticulo, distancia_objetivo_m, incluir_borde)

    # 3) Gráficos principales
    graficar_poligono_y_origen(poligono, origen_reticulo, nombre_fig_origen, 'Polígono y punto inicial (' + str(int(distancia_objetivo_m)) + ' m)')
    graficar_poligono_y_puntos(poligono, origen_reticulo, puntos, nombre_fig_puntos, 'Puntos dentro del polígono (d = ' + str(int(distancia_objetivo_m)) + ' m)')

    # 4) Exportaciones
    exportar_csv_puntos(puntos, nombre_csv_puntos)
    exportar_csv_vertices(vertices, nombre_csv_vertices)
    exportar_dxf(poligono, puntos, nombre_dxf)

    # 5) Informe Word
    crear_informe_word(poligono, vertices, origen_reticulo, puntos, nombre_fig_origen, nombre_fig_puntos, nombre_informe, archivo_excel, distancia_objetivo_m)

    # 6) Análisis de puntos problemáticos según criterio solicitado
    df_problematicos = analizar_puntos_problematicos(poligono, puntos, distancia_objetivo_m, tolerancia_m, nombre_csv_problematicos, nombre_fig_problematicos)

    # 7) Mensajes finales
    print('CSV puntos: ' + nombre_csv_puntos)
    print('CSV vértices: ' + nombre_csv_vertices)
    if DISPONIBLE_EZDXF:
        print('DXF: ' + nombre_dxf)
    else:
        print('DXF omitido (instala ezdxf para exportar)')

    if DISPONIBLE_DOCX:
        print('Informe Word: ' + nombre_informe)
    else:
        print('Informe Word omitido (instala python-docx para exportar)')

    print('Total de puntos generados: ' + str(len(puntos)))
    print('CSV puntos problemáticos: ' + nombre_csv_problematicos)
    if df_problematicos is not None:
        print('Total problemáticos: ' + str(df_problematicos.shape[0]))