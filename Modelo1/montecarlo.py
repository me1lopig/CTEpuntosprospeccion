# Monte Carlo para ajustar malla hexagonal (offset + rotación) maximizando puntos dentro del polígono.
# Luego, detección de puntos problemáticos (sin vecino a distancia objetivo y/o fuera de la componente principal).
# Requiere: pandas, numpy, shapely, matplotlib. Opcional: ezdxf, python-docx, scipy.

import os
import math
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from shapely.geometry import Polygon, Point, LineString

# ---------------------------
# Parámetros
# ---------------------------
excel_path = 'poligono.xlsx'     # Debe tener columnas X e Y
sheet_name = 0
distancia_m = 45.0               # Distancia objetivo de la malla
incluir_borde = True             # covers en vez de contains
ensayos_mc = 1000                # Número de pruebas Monte Carlo
semilla = 7                      # Reproducibilidad
tol_dist = 0.5                   # Tolerancia para enlaces a distancia_m

# Archivos de salida
csv_puntos_mc = 'puntos_hex_mc_' + str(int(distancia_m)) + 'm.csv'
csv_vertices = 'poligono_vertices.csv'
dxf_mc = 'reticulo_mc_' + str(int(distancia_m)) + 'm.dxf'
doc_mc = 'informe_mc_' + str(int(distancia_m)) + 'm.docx'
fig_mc_origen = 'mc_mejor_origen.png'
fig_mc_puntos = 'mc_mejor_puntos.png'
csv_problem = 'puntos_problematicos_mc_' + str(int(distancia_m)) + 'm.csv'
fig_problem = 'puntos_problematicos_mc_' + str(int(distancia_m)) + 'm.png'

# ---------------------------
# Utilidades
# ---------------------------
def cargar_poligono_desde_excel(path_excel, hoja):
    df = pd.read_excel(path_excel, sheet_name=hoja)
    df.columns = [str(c).strip() for c in df.columns]
    low = [c.lower() for c in df.columns]
    colx = df.columns[low.index('x')]
    coly = df.columns[low.index('y')]
    df_v = df[[colx, coly]].dropna().copy()
    df_v[colx] = pd.to_numeric(df_v[colx], errors='coerce')
    df_v[coly] = pd.to_numeric(df_v[coly], errors='coerce')
    df_v = df_v.dropna()
    poly = Polygon(list(df_v.itertuples(index=False, name=None)))
    if not poly.is_valid:
        poly = poly.buffer(0)
    df_v.columns = ['X','Y']
    return df_v, poly

def generar_hex_rotado(poligono, origen, distancia, theta_rad, incluir_borde=True):
    c = np.cos(theta_rad)
    s = np.sin(theta_rad)
    vstep = distancia * np.sqrt(3) / 2.0
    verts_np = np.array(poligono.exterior.coords)
    radio = float(np.max(np.linalg.norm(verts_np[:, :2] - np.array(origen), axis=1))) + distancia * 2.0
    filas = int(np.ceil(radio / vstep)) + 2
    cols = int(np.ceil(radio / distancia)) + 2
    pts = []
    for i in range(-filas, filas + 1):
        for j in range(-cols, cols + 1):
            bx = distancia * (j + (0.5 if (i % 2) != 0 else 0.0))
            by = vstep * i
            x = origen[0] + c * bx - s * by
            y = origen[1] + s * bx + c * by
            ok = poligono.covers(Point(x, y)) if incluir_borde else poligono.contains(Point(x, y))
            if ok:
                pts.append((x, y))
    return pts

def monte_carlo_hex(poligono, distancia, incluir_borde, n_trials, seed):
    np.random.seed(seed)
    centroid = tuple(poligono.centroid.coords)[0]
    vstep = distancia * np.sqrt(3) / 2.0
    best_count = -1
    best_theta = 0.0
    best_off = (0.0, 0.0)
    best_pts = []
    for k in range(n_trials):
        theta = np.random.uniform(0.0, np.pi / 3.0)  # simetría 60°
        ox = np.random.uniform(0.0, distancia)
        oy = np.random.uniform(0.0, vstep)
        origen_try = (centroid[0] + ox, centroid[1] + oy)
        pts_try = generar_hex_rotado(poligono, origen_try, distancia, theta, incluir_borde)
        if len(pts_try) > best_count:
            best_count = len(pts_try)
            best_theta = theta
            best_off = (ox, oy)
            best_pts = pts_try
    return best_pts, best_theta, best_off, centroid, best_count

def exportar_dxf(poligono, puntos, ruta_dxf):
    try:
        import ezdxf
        doc = ezdxf.new(setup=True)
        ms = doc.modelspace()
        coords = list(poligono.exterior.coords)
        if coords[0] != coords[-1]:
            coords.append(coords[0])
        ms.add_lwpolyline(coords, format='xy', dxfattribs={'closed': True})
        for x, y in puntos:
            ms.add_point((x, y))
        doc.saveas(ruta_dxf)
        return True
    except Exception:
        return False

def exportar_doc(fig_origen, fig_puntos, ruta_doc, trials, theta, off, count, distancia):
    try:
        from docx import Document
        from docx.shared import Cm
        rep = Document()
        rep.add_heading('Informe Monte Carlo malla hexagonal (d = ' + str(int(distancia)) + ' m)', level=1)
        rep.add_paragraph('Ensayos Monte Carlo: ' + str(trials))
        rep.add_paragraph('Ángulo óptimo (grados): ' + str(round(theta * 180.0 / np.pi, 2)))
        rep.add_paragraph('Offset óptimo: ' + str(round(off[0],2)) + ', ' + str(round(off[1],2)))
        rep.add_paragraph('Total de puntos dentro: ' + str(count))
        rep.add_paragraph('Polígono y origen:')
        rep.add_picture(fig_origen, width=Cm(14))
        rep.add_paragraph('Puntos dentro del polígono:')
        rep.add_picture(fig_puntos, width=Cm(14))
        rep.save(ruta_doc)
        return True
    except Exception:
        return False

def graficar_mc(poligono, centroid, off, puntos, fig1, fig2, count, distancia):
    xp, yp = poligono.exterior.xy
    # Origen
    plt.figure(figsize=(6,5))
    plt.fill(xp, yp, alpha=0.4, fc='lightblue', ec='darkblue')
    plt.scatter([centroid[0] + off[0]], [centroid[1] + off[1]], c='red', s=60, label='Origen')
    plt.title('Mejor Monte Carlo (puntos dentro: ' + str(count) + ')')
    plt.xlabel('X'); plt.ylabel('Y'); plt.axis('equal'); plt.grid(True); plt.legend()
    plt.savefig(fig1, dpi=180, bbox_inches='tight')
    plt.show()
    # Puntos
    plt.figure(figsize=(6,5))
    plt.fill(xp, yp, alpha=0.35, fc='lightblue', ec='darkblue')
    if len(puntos) > 0:
        arr = np.array(puntos)
        plt.scatter(arr[:,0], arr[:,1], s=16, c='green', label='Puntos')
    plt.scatter([centroid[0] + off[0]], [centroid[1] + off[1]], c='red', s=60, label='Origen')
    plt.title('Monte Carlo: malla óptima (d = ' + str(int(distancia)) + ' m)')
    plt.xlabel('X'); plt.ylabel('Y'); plt.axis('equal'); plt.grid(True); plt.legend()
    plt.savefig(fig2, dpi=180, bbox_inches='tight')
    plt.show()

def detectar_puntos_problematicos(poligono, puntos, distancia, tol, ruta_csv, ruta_fig):
    if len(puntos) == 0:
        pd.DataFrame(columns=['X','Y','sin_vecino','fuera_componente_principal']).to_csv(ruta_csv, index=False)
        return pd.DataFrame()
    pts_arr = np.array(puntos)
    # Vecinos candidatos por KDTree si está disponible
    try:
        from scipy.spatial import cKDTree as KDTree
        tree = KDTree(pts_arr)
        cand = tree.query_ball_point(pts_arr, distancia + tol)
    except Exception:
        cand = []
        for i in range(len(pts_arr)):
            d = np.linalg.norm(pts_arr - pts_arr[i], axis=1)
            vecinos = list(np.where((d <= distancia + tol) & (d >= distancia - tol))[0])
            cand.append(vecinos)
    # Grafo de aristas válidas (segmento dentro del polígono)
    n = len(pts_arr)
    adj = [[] for _ in range(n)]
    for i in range(n):
        for j in cand[i]:
            if j == i:
                continue
            dij = np.linalg.norm(pts_arr[i] - pts_arr[j])
            if dij < distancia - tol or dij > distancia + tol:
                continue
            seg = LineString([tuple(pts_arr[i]), tuple(pts_arr[j])])
            if poligono.covers(seg):
                adj[i].append(j)
    sin_vecino = [i for i in range(n) if len(adj[i]) == 0]
    # Componentes
    visit = np.zeros(n, dtype=bool)
    comps = []
    for i in range(n):
        if not visit[i]:
            cola = [i]
            visit[i] = True
            comp = [i]
            while cola:
                u = cola.pop(0)
                for v in adj[u]:
                    if not visit[v]:
                        visit[v] = True
                        cola.append(v)
                        comp.append(v)
            comps.append(comp)
    main = set(max(comps, key=len)) if len(comps) > 0 else set()
    fuera_main = [i for i in range(n) if i not in main]
    # Resultado
    mask = np.zeros(n, dtype=bool)
    for i in sin_vecino:
        mask[i] = True
    for i in fuera_main:
        mask[i] = True
    df_problem = pd.DataFrame({
        'X': pts_arr[:,0],
        'Y': pts_arr[:,1],
        'sin_vecino': [i in sin_vecino for i in range(n)],
        'fuera_componente_principal': [i in fuera_main for i in range(n)]
    })
    df_problem_only = df_problem[mask].copy()
    df_problem_only.to_csv(ruta_csv, index=False)
    # Figura
    xp, yp = poligono.exterior.xy
    plt.figure(figsize=(6,5))
    plt.fill(xp, yp, alpha=0.35, fc='lightblue', ec='darkblue')
    plt.scatter(pts_arr[:,0], pts_arr[:,1], s=18, c='gray', label='Puntos')
    if df_problem_only.shape[0] > 0:
        plt.scatter(df_problem_only['X'], df_problem_only['Y'], s=45, c='red', label='Problemáticos')
    plt.title('Puntos problemáticos a ' + str(int(distancia)) + ' m (MC)')
    plt.xlabel('X'); plt.ylabel('Y'); plt.axis('equal'); plt.grid(True); plt.legend()
    plt.savefig(ruta_fig, dpi=180, bbox_inches='tight')
    plt.show()
    return df_problem_only

# ---------------------------
# Flujo principal
# ---------------------------
df_vertices, poligono = cargar_poligono_desde_excel(excel_path, sheet_name)
puntos_mc, theta_opt, off_opt, centroide, best_count = monte_carlo_hex(
    poligono, distancia_m, incluir_borde, ensayos_mc, semilla
)

# Gráficos MC
graficar_mc(poligono, centroide, off_opt, puntos_mc, fig_mc_origen, fig_mc_puntos, best_count, distancia_m)

# Exportaciones
pd.DataFrame(puntos_mc, columns=['X','Y']).to_csv(csv_puntos_mc, index=False)
df_vertices.to_csv(csv_vertices, index=False)
ok_dxf = exportar_dxf(poligono, puntos_mc, dxf_mc)
ok_doc = exportar_doc(fig_mc_origen, fig_mc_puntos, doc_mc, ensayos_mc, theta_opt, off_opt, best_count, distancia_m)

print('CSV puntos MC:', csv_puntos_mc)
print('CSV vértices:', csv_vertices)
print('DXF:', dxf_mc if ok_dxf else 'DXF omitido (instala ezdxf)')
print('Informe Word:', doc_mc if ok_doc else 'Informe omitido (instala python-docx)')
print('Puntos dentro (MC):', best_count)
print('Ángulo óptimo (rad):', theta_opt)
print('Offset óptimo:', off_opt)

# Análisis de problemáticos
df_prob = detectar_puntos_problematicos(poligono, puntos_mc, distancia_m, tol_dist, csv_problem, fig_problem)
print('CSV problemáticos:', csv_problem)
print('Total problemáticos:', 0 if df_prob is None else df_prob.shape[0])