import streamlit as st
import zipfile
import xml.etree.ElementTree as ET
import numpy as np
import matplotlib.pyplot as plt
import math
import utm
import pandas as pd
from shapely.geometry import Polygon as ShapelyPolygon, Point as ShapelyPoint
import contextily as ctx
import io
import tempfile
from docx import Document
from docx.shared import Inches

st.set_page_config(page_title="Generador de Malla Agrícola/Topográfica", layout="wide")
st.title("📍 Generador de Puntos sobre Polígono (KMZ)")

# --- FUNCIONES DE MEMORIA ---
def limpiar_descargas():
    if 'archivos_listos' in st.session_state:
        del st.session_state['archivos_listos']
        del st.session_state['excel_data']
        del st.session_state['word_data']

# --- FUNCIONES DE EXPORTACIÓN ---
def generar_excel(df_poligono, df_puntos):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df_poligono.to_excel(writer, sheet_name='Poligono_Original', index=False)
        df_puntos.to_excel(writer, sheet_name='Puntos_Generados', index=False)
    return output.getvalue()

def generar_informe_word(area_ha, area_m2, num_puntos, area_por_punto, puntos_por_ha, distancia, margen, metodo, angulo_opt, fig):
    doc = Document()
    doc.add_heading('INFORME DE REPLANTEO Y MALLA DE PUNTOS', 0)
    
    doc.add_heading('1. Resumen de Superficie y Densidad:', level=1)
    doc.add_paragraph(f"• Superficie del polígono: {area_ha:.2f} hectáreas ({area_m2:,.0f} m²)")
    doc.add_paragraph(f"• Total de puntos generados: {num_puntos}")
    doc.add_paragraph(f"• Relación de densidad: {puntos_por_ha:.0f} puntos/ha ({area_por_punto:.2f} m²/punto)")

    doc.add_heading('2. Parámetros de Configuración:', level=1)
    doc.add_paragraph(f"• Método de distribución: {metodo}")
    doc.add_paragraph(f"• Distancia entre puntos: {distancia} metros")
    doc.add_paragraph(f"• Distancia de seguridad al borde (margen): {margen} metros")
    if "OPTIMIZADO" in metodo:
        doc.add_paragraph(f"• Ángulo óptimo calculado: {angulo_opt} grados")

    doc.add_heading('3. Plano de Distribución:', level=1)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp_img:
        fig.savefig(tmp_img.name, format="png", bbox_inches="tight", dpi=300)
        doc.add_picture(tmp_img.name, width=Inches(6.0))
        
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

# --- PROCESAMIENTO PRINCIPAL ---
def procesar_kmz(uploaded_file, distancia_metros, distancia_borde_metros, tipo_mapa, opacidad, metodo):
    with zipfile.ZipFile(uploaded_file, 'r') as kmz:
        kml_filename = [f for f in kmz.namelist() if f.endswith('.kml')][0]
        with kmz.open(kml_filename, 'r') as kml_file:
            root = ET.parse(kml_file).getroot()

    ns = {'kml': root.tag.split('}')[0].strip('{')} if '}' in root.tag else {}
    polygon = root.find('.//kml:Polygon', ns) if ns else root.find('.//Polygon')
    coords_text = (polygon.find('.//kml:outerBoundaryIs/kml:LinearRing/kml:coordinates', ns) if ns 
                   else polygon.find('.//outerBoundaryIs/LinearRing/coordinates')).text.strip()
    coords_geo = np.array([(float(p.split(',')[0]), float(p.split(',')[1])) for p in coords_text.split()])

    datos_poligono = []
    for i, (lon, lat) in enumerate(coords_geo):
        e, n, z_num, z_let = utm.from_latlon(lat, lon)
        datos_poligono.append({
            'Vertice': i + 1, 'Latitud': lat, 'Longitud': lon,
            'UTM_X': round(e, 2), 'UTM_Y': round(n, 2), 'Huso': f"{z_num}{z_let}"
        })
    df_poligono = pd.DataFrame(datos_poligono)

    lat_media = np.mean(coords_geo[:, 1])
    lat_rad = np.radians(lat_media)
    m_por_deg_lat = 111132.92 - 559.82 * np.cos(2 * lat_rad)
    m_por_deg_lon = 111412.84 * np.cos(lat_rad) - 93.5 * np.cos(3 * lat_rad)
    
    min_lon, min_lat = np.min(coords_geo, axis=0)
    max_lon, max_lat = np.max(coords_geo, axis=0)
    poly_x = (coords_geo[:, 0] - min_lon) * m_por_deg_lon
    poly_y = (coords_geo[:, 1] - min_lat) * m_por_deg_lat
    poly_m = np.column_stack((poly_x, poly_y))

    poligono_base = ShapelyPolygon(poly_m)
    area_m2 = poligono_base.area

    # Aplicar margen de seguridad
    if distancia_borde_metros > 0:
        poligono_util = poligono_base.buffer(-distancia_borde_metros)
    else:
        poligono_util = poligono_base

    if poligono_util.is_empty:
        return None, None, None, "El margen de seguridad es demasiado grande.", area_m2, 0

    # Centro y radio para rotaciones amplias
    cx, cy = np.mean(poly_m[:, 0]), np.mean(poly_m[:, 1])
    R = math.hypot(max_lon - min_lon, max_lat - min_lat) * max(m_por_deg_lon, m_por_deg_lat)
    
    dy = distancia_metros * math.sin(math.pi/3)
    dx = distancia_metros
    
    pasos_y = int((2 * R) / dy) + 4
    pasos_x = int((2 * R) / dx) + 4
    start_x, start_y = cx - R, cy - R
    
    # Crear malla base gigante centrada
    puntos_base = []
    for fila in range(pasos_y):
        y = start_y + fila * dy
        offset_x = (distancia_metros / 2) if fila % 2 == 1 else 0
        for col in range(pasos_x):
            x = start_x + col * dx + offset_x
            puntos_base.append([x, y])
    puntos_base = np.array(puntos_base)

    mejor_angulo = 0
    mejores_puntos_finales = []
    max_puntos_dentro = -1

    # --- AQUÍ OCURRE LA MAGIA DEL "BOMBARDEO" ---
    angulos_a_probar = range(0, 60, 1) if "OPTIMIZADO" in metodo else [0]
    
    for angulo in angulos_a_probar:
        # Rotar la malla entera
        angulo_rad = math.radians(angulo)
        cos_a, sin_a = math.cos(angulo_rad), math.sin(angulo_rad)
        
        nx = cx + (puntos_base[:, 0] - cx) * cos_a - (puntos_base[:, 1] - cy) * sin_a
        ny = cy + (puntos_base[:, 0] - cx) * sin_a + (puntos_base[:, 1] - cy) * cos_a
        puntos_rotados = np.column_stack((nx, ny))
        
        # Filtrar cuántos caen dentro
        mascara = np.array([poligono_util.contains(ShapelyPoint(pt[0], pt[1])) for pt in puntos_rotados])
        puntos_validos = puntos_rotados[mascara]
        
        # Si este ángulo mete más puntos, lo guardamos como el nuevo rey
        if len(puntos_validos) > max_puntos_dentro:
            max_puntos_dentro = len(puntos_validos)
            mejores_puntos_finales = puntos_validos
            mejor_angulo = angulo

    if len(mejores_puntos_finales) == 0:
        return None, None, None, "La distancia o el margen es demasiado grande. No cupo ningún punto.", area_m2, 0

    # Volver a geográficas
    puntos_finales_lon = (mejores_puntos_finales[:, 0] / m_por_deg_lon) + min_lon
    puntos_finales_lat = (mejores_puntos_finales[:, 1] / m_por_deg_lat) + min_lat

    datos_exportar = []
    for lon, lat in zip(puntos_finales_lon, puntos_finales_lat):
        easting, northing, zone_number, zone_letter = utm.from_latlon(lat, lon)
        datos_exportar.append({
            'ID_Punto': len(datos_exportar) + 1, 'Latitud': lat, 'Longitud': lon,
            'UTM_X': round(easting, 2), 'UTM_Y': round(northing, 2), 'Huso': f"{zone_number}{zone_letter}"
        })
    df_puntos = pd.DataFrame(datos_exportar)

    # --- CREAR GRÁFICO ---
    fig, ax = plt.subplots(figsize=(10, 8))
    borde_x = df_poligono['UTM_X'].tolist()
    borde_y = df_poligono['UTM_Y'].tolist()
        
    ax.plot(borde_x, borde_y, 'k-', linewidth=2, label='Límite Real')

    if distancia_borde_metros > 0 and poligono_util.geom_type == 'Polygon':
        x_inner, y_inner = poligono_util.exterior.xy
        inner_lon = (np.array(x_inner) / m_por_deg_lon) + min_lon
        inner_lat = (np.array(y_inner) / m_por_deg_lat) + min_lat
        
        inner_x, inner_y = [], []
        for lon, lat in zip(inner_lon, inner_lat):
            e, n, _, _ = utm.from_latlon(lat, lon)
            inner_x.append(e)
            inner_y.append(n)
            
        ax.plot(inner_x, inner_y, 'b--', linewidth=1.5, label=f'Margen ({distancia_borde_metros}m)')
        ax.fill(inner_x, inner_y, alpha=opacidad, color='cyan')
    else:
        ax.fill(borde_x, borde_y, alpha=opacidad, color='cyan')

    ax.scatter(df_puntos['UTM_X'], df_puntos['UTM_Y'], c='red', s=15, label='Puntos', zorder=5)
    
    titulo_mapa = f"Plano de Replanteo a {distancia_metros}m"
    if "OPTIMIZADO" in metodo:
        titulo_mapa += f"\n(Optimizado: Rotado {mejor_angulo}º para máxima densidad)"
        
    ax.set_title(titulo_mapa)
    ax.set_xlabel("UTM Este (X) [metros]")
    ax.set_ylabel("UTM Norte (Y) [metros]")
    ax.ticklabel_format(useOffset=False, style='plain')
    plt.xticks(rotation=15)
    ax.axis('equal')
    ax.legend()
    ax.grid(True, linestyle=':', alpha=0.6)

    if tipo_mapa != "Ninguno":
        _, _, huso_num, _ = utm.from_latlon(lat_media, min_lon)
        epsg_code = 32600 + huso_num 
        url_pnoa = "https://www.ign.es/wmts/pnoa-ma?request=GetTile&service=WMTS&version=1.0.0&Layer=OI.OrthoimageCoverage&Style=default&Format=image/jpeg&TileMatrixSet=GoogleMapsCompatible&TileMatrix={z}&TileRow={y}&TileCol={x}"
        url_mtn = "https://www.ign.es/wmts/mapa-raster?request=getTile&layer=MTN&TileMatrixSet=GoogleMapsCompatible&TileMatrix={z}&TileCol={x}&TileRow={y}&format=image/jpeg"
        fuente = url_pnoa if tipo_mapa == "PNOA (Satélite)" else url_mtn
        try:
            ctx.add_basemap(ax, crs=f"EPSG:{epsg_code}", source=fuente, alpha=1.0)
        except Exception:
            pass

    return df_poligono, df_puntos, fig, None, area_m2, mejor_angulo

# --- INTERFAZ DE USUARIO ---
col1, col2 = st.columns([1, 2])

with col1:
    st.header("1. Configuración")
    archivo_kmz = st.file_uploader("Sube tu archivo .kmz", type=['kmz'], on_change=limpiar_descargas)
    
    st.markdown("### Geometría de Puntos")
    metodo_dist = st.selectbox("Método de Distribución:", [
        "Hexagonal Normal (Norte-Sur)",
        "Hexagonal OPTIMIZADO (Búsqueda del Máximo)"
    ], on_change=limpiar_descargas)
    
    distancia = st.number_input("Distancia mínima entre puntos (metros):", min_value=1.0, value=30.0, step=1.0, on_change=limpiar_descargas)
    margen = st.number_input("Distancia al borde (metros):", min_value=0.0, value=0.0, step=1.0, on_change=limpiar_descargas)
    
    st.markdown("---")
    st.subheader("Visualización")
    tipo_mapa = st.radio("Fondo de Mapa (IGN):", ["PNOA (Satélite)", "MTN (Topográfico)", "Ninguno"], on_change=limpiar_descargas)
    transparencia = st.slider("Transparencia del polígono", min_value=0.0, max_value=1.0, value=0.3, step=0.1, on_change=limpiar_descargas)
    
with col2:
    st.header("2. Resultados")
    if archivo_kmz is not None:
        texto_carga = 'Escaneando ángulos para máxima densidad...' if 'OPTIMIZADO' in metodo_dist else 'Procesando geometría...'
        with st.spinner(texto_carga):
            df_poli, df_pts, figura, error_msg, area_m2, angulo_opt = procesar_kmz(archivo_kmz, distancia, margen, tipo_mapa, transparencia, metodo_dist)
            
            if error_msg:
                st.error(error_msg)
            else:
                num_puntos = len(df_pts)
                area_ha = area_m2 / 10000 
                area_por_punto = area_m2 / num_puntos if num_puntos > 0 else 0
                puntos_por_ha = num_puntos / area_ha if area_ha > 0 else 0

                st.subheader("📊 Resumen de la Parcela")
                m1, m2, m3 = st.columns(3)
                m1.metric(label="Área Total", value=f"{area_ha:.2f} ha", delta=f"{area_m2:,.0f} m²", delta_color="off")
                
                # Destacamos si se ha optimizado
                if "OPTIMIZADO" in metodo_dist:
                    m2.metric(label="Total de Puntos", value=f"{num_puntos} pts", delta=f"Rotado {angulo_opt}º", delta_color="normal")
                else:
                    m2.metric(label="Total de Puntos", value=f"{num_puntos} pts")
                    
                m3.metric(label="Densidad", value=f"{area_por_punto:.1f} m²/punto", delta=f"{puntos_por_ha:.0f} pts/ha", delta_color="off")

                st.pyplot(figura)
                
                st.markdown("---")
                st.markdown("### 📥 Exportar Resultados")
                
                if not st.session_state.get('archivos_listos', False):
                    if st.button("🚀 Finalizar y Preparar Informes", type="primary"):
                        with st.spinner("Generando Word y Excel..."):
                            st.session_state['excel_data'] = generar_excel(df_poli, df_pts)
                            st.session_state['word_data'] = generar_informe_word(area_ha, area_m2, num_puntos, area_por_punto, puntos_por_ha, distancia, margen, metodo_dist, angulo_opt, figura)
                            st.session_state['archivos_listos'] = True
                            st.rerun()
                
                if st.session_state.get('archivos_listos', False):
                    st.success("✅ ¡Archivos generados con éxito!")
                    col_btn1, col_btn2 = st.columns(2)
                    col_btn1.download_button(label="📊 Descargar Excel", data=st.session_state['excel_data'], file_name=f'Replanteo_{int(distancia)}m.xlsx', mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
                    col_btn2.download_button(label="📝 Descargar Word", data=st.session_state['word_data'], file_name=f'Informe_Topografico_{int(distancia)}m.docx', mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')