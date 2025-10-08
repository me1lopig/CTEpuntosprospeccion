# Ejecutar el script con variables en español: lee poligono.xlsx, genera retículo hexagonal (30 m),
# optimiza offset, exporta CSV y DXF, y crea el informe Word con imágenes en centímetros.

import pandas as pd
import numpy as np
from shapely.geometry import Polygon, Point
import matplotlib.pyplot as plt
import ezdxf
from docx import Document
from docx.shared import Cm

# Configuración
archivo_excel = 'poligono.xlsx'
hoja_excel = 0
separacion_m = 45.0
incluir_borde = True
optimizar_desplazamiento = True

# Carga y limpieza
df_poligono = pd.read_excel(archivo_excel, sheet_name=hoja_excel)
df_poligono.columns = [str(c).strip() for c in df_poligono.columns]
columnas_minus = [c.lower() for c in df_poligono.columns]
columna_x = df_poligono.columns[columnas_minus.index('x')]
columna_y = df_poligono.columns[columnas_minus.index('y')]

x_numerico = pd.to_numeric(df_poligono[columna_x], errors='coerce')
y_numerico = pd.to_numeric(df_poligono[columna_y], errors='coerce')
filtro_numerico = x_numerico.notna() & y_numerico.notna()
df_coordenadas = df_poligono.loc[filtro_numerico, [columna_x, columna_y]].copy()
df_coordenadas[columna_x] = pd.to_numeric(df_coordenadas[columna_x], errors='coerce')
df_coordenadas[columna_y] = pd.to_numeric(df_coordenadas[columna_y], errors='coerce')

vertices = list(df_coordenadas.itertuples(index=False, name=None))
poligono = Polygon(vertices)
if not poligono.is_valid:
    poligono = poligono.buffer(0)

centroide = tuple(poligono.centroid.coords)[0]

# Generación del retículo
def generar_puntos_con_centro(poligono, centro, distancia, incluir_borde=True):
    verts_np = np.array(poligono.exterior.coords)
    radio = float(np.max(np.linalg.norm(verts_np[:, :2] - np.array(centro), axis=1)))
    filas = int(np.ceil(radio / (distancia * np.sqrt(3) / 2.0))) + 2
    puntos_locales = []
    for i in range(-filas, filas + 1):
        for j in range(-filas, filas + 1):
            x = centro[0] + distancia * (j + (0.5 if (i % 2) != 0 else 0.0))
            y = centro[1] + (distancia * np.sqrt(3) / 2.0) * i
            adentro = poligono.covers(Point(x, y)) if incluir_borde else poligono.contains(Point(x, y))
            if adentro:
                puntos_locales.append((x, y))
    return puntos_locales

# Optimización del desplazamiento
paso_vertical = separacion_m * np.sqrt(3) / 2.0
if optimizar_desplazamiento:
    xs = np.linspace(0.0, separacion_m, 21)
    ys = np.linspace(0.0, paso_vertical, 21)
    mejor_cantidad = -1
    mejor_despl = (0.0, 0.0)
    mejores_puntos = []
    for oy in ys:
        for ox in xs:
            c = (centroide[0] + ox, centroide[1] + oy)
            pts_tmp = generar_puntos_con_centro(poligono, c, separacion_m, incluir_borde)
            if len(pts_tmp) > mejor_cantidad:
                mejor_cantidad = len(pts_tmp)
                mejor_despl = (ox, oy)
                mejores_puntos = pts_tmp
    origen_reticulo = (centroide[0] + mejor_despl[0], centroide[1] + mejor_despl[1])
    puntos = mejores_puntos
else:
    origen_reticulo = centroide
    puntos = generar_puntos_con_centro(poligono, origen_reticulo, separacion_m, incluir_borde)

# Gráficos y guardado de imágenes
xp, yp = poligono.exterior.xy
plt.figure(figsize=(6,5))
plt.fill(xp, yp, alpha=0.4, fc='lightblue', ec='darkblue')
plt.scatter(origen_reticulo[0], origen_reticulo[1], c='red', s=60, label='Punto inicial')
plt.title('Polígono y punto inicial')
plt.xlabel('X')
plt.ylabel('Y')
plt.axis('equal')
plt.grid(True)
plt.legend()
ruta_figura_origen = 'fig_poligono_origen.png'
plt.savefig(ruta_figura_origen, dpi=180, bbox_inches='tight')
plt.show()

plt.figure(figsize=(6,5))
plt.fill(xp, yp, alpha=0.4, fc='lightblue', ec='darkblue')
if len(puntos) > 0:
    arr = np.array(puntos)
    plt.scatter(arr[:, 0], arr[:, 1], c='green', s=15, label='Puntos dentro')
plt.scatter(origen_reticulo[0], origen_reticulo[1], c='red', s=60, label='Punto inicial')
plt.title('Puntos dentro del polígono d = ' + str(int(separacion_m)) + ' m')
plt.xlabel('X')
plt.ylabel('Y')
plt.axis('equal')
plt.grid(True)
plt.legend()
ruta_figura_puntos = 'fig_poligono_puntos.png'
plt.savefig(ruta_figura_puntos, dpi=180, bbox_inches='tight')
plt.show()

# Exportaciones CSV y DXF
nombre_csv_puntos = 'puntos_hex_' + str(int(separacion_m)) + 'm.csv'
nombre_csv_vertices = 'poligono_vertices.csv'
nombre_dxf = 'reticulo_poligono_' + str(int(separacion_m)) + 'm.dxf'

pd.DataFrame(puntos, columns=['X', 'Y']).to_csv(nombre_csv_puntos, index=False)
pd.DataFrame(vertices, columns=['X', 'Y']).to_csv(nombre_csv_vertices, index=False)

# DXF
archivo_dxf = ezdxf.new(setup=True)
espacio_modelo = archivo_dxf.modelspace()
coords_polilinea = list(vertices)
if coords_polilinea[0] != coords_polilinea[-1]:
    coords_polilinea.append(coords_polilinea[0])
espacio_modelo.add_lwpolyline(coords_polilinea, format='xy', dxfattribs={'closed': True})
for x, y in puntos:
    espacio_modelo.add_point((x, y))
archivo_dxf.saveas(nombre_dxf)

# Informe Word
nombre_informe = 'informe_reticulo_hex_' + str(int(separacion_m)) + 'm.docx'

documento = Document()
documento.add_heading('Informe de retículo hexagonal en polígono', level=1)

documento.add_heading('Datos de entrada', level=2)
documento.add_paragraph('Archivo: ' + archivo_excel)
documento.add_paragraph('Vértices del polígono (X, Y):')
tabla_vertices = documento.add_table(rows=1, cols=2)
tabla_vertices.rows[0].cells[0].text = 'X'
tabla_vertices.rows[0].cells[1].text = 'Y'
for x, y in vertices:
    fila = tabla_vertices.add_row().cells
    fila[0].text = str(x)
    fila[1].text = str(y)

documento.add_paragraph('Punto inicial del retículo (X, Y): ' + str(round(origen_reticulo[0], 3)) + ', ' + str(round(origen_reticulo[1], 3)))
documento.add_paragraph('Polígono y punto inicial:')
documento.add_picture(ruta_figura_origen, width=Cm(14))

# Resultados
documento.add_heading('Resultados', level=2)
documento.add_paragraph('Separación del retículo: ' + str(int(separacion_m)) + ' m')
documento.add_paragraph('Número de puntos dentro del polígono (incluyendo el inicial): ' + str(len(puntos)))
documento.add_paragraph('Polígono con los puntos dentro (punto inicial en rojo):')
documento.add_picture(ruta_figura_puntos, width=Cm(14))

# Guardar
documento.save(nombre_informe)

print('Listo: CSV, DXF e informe actualizados')
print('CSV puntos: ' + nombre_csv_puntos)
print('CSV polígono: ' + nombre_csv_vertices)
print('DXF: ' + nombre_dxf)
print('Informe: ' + nombre_informe)
print('Total puntos: ' + str(len(puntos)))