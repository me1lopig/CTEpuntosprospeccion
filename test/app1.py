# Lee poligono.xlsx, genera retículo hexagonal (30 m) optimizado, grafica,
# exporta CSV + DXF y crea un informe Word con tablas e imágenes.

import pandas as pd
import numpy as np
from shapely.geometry import Polygon, Point
import matplotlib.pyplot as plt
import ezdxf
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# -----------------------
# Configuración
# -----------------------
excel_filename = 'poligono.xlsx'
sheet_name = 0
spacing_m = 30.0 # separación de los puntos en metros
include_border = True
optimize_offsets = True

# -----------------------
# Carga y limpieza
# -----------------------
df = pd.read_excel(excel_filename, sheet_name=sheet_name)
df.columns = [str(c).strip() for c in df.columns]
cols_lower = [c.lower() for c in df.columns]
if 'x' in cols_lower and 'y' in cols_lower:
    col_x = df.columns[cols_lower.index('x')]
    col_y = df.columns[cols_lower.index('y')]
else:
    raise ValueError('No se encontraron columnas X/Y en el Excel')

x_num = pd.to_numeric(df[col_x], errors='coerce')
y_num = pd.to_numeric(df[col_y], errors='coerce')
mask_numeric = x_num.notna() & y_num.notna()
coords_df = df.loc[mask_numeric, [col_x, col_y]].copy()
coords_df[col_x] = pd.to_numeric(coords_df[col_x], errors='coerce')
coords_df[col_y] = pd.to_numeric(coords_df[col_y], errors='coerce')

vertices = list(coords_df.itertuples(index=False, name=None))
if len(vertices) < 3:
    raise ValueError('Se requieren al menos 3 vértices para el polígono')

poly = Polygon(vertices)
if not poly.is_valid:
    poly = poly.buffer(0)

centroid = tuple(poly.centroid.coords)[0]

# -----------------------
# Generación de retículo
# -----------------------
def generate_points_with_center(poly, c, d, include_border=True):
    verts = np.array(poly.exterior.coords)
    radius = float(np.max(np.linalg.norm(verts[:, :2] - np.array(c), axis=1)))
    n = int(np.ceil(radius / (d * np.sqrt(3) / 2.0))) + 2
    pts = []
    for i in range(-n, n + 1):
        for j in range(-n, n + 1):
            x = c[0] + d * (j + (0.5 if (i % 2) != 0 else 0.0))
            y = c[1] + (d * np.sqrt(3) / 2.0) * i
            inside = poly.covers(Point(x, y)) if include_border else poly.contains(Point(x, y))
            if inside:
                pts.append((x, y))
    return pts

step_v = spacing_m * np.sqrt(3) / 2.0
if optimize_offsets:
    xs = np.linspace(0.0, spacing_m, 21)
    ys = np.linspace(0.0, step_v, 21)
    best_count = -1
    best_off = (0.0, 0.0)
    best_pts = []
    for oy in ys:
        for ox in xs:
            c = (centroid[0] + ox, centroid[1] + oy)
            pts_tmp = generate_points_with_center(poly, c, spacing_m, include_border)
            if len(pts_tmp) > best_count:
                best_count = len(pts_tmp)
                best_off = (ox, oy)
                best_pts = pts_tmp
    origin = (centroid[0] + best_off[0], centroid[1] + best_off[1])
    puntos = best_pts
else:
    origin = centroid
    puntos = generate_points_with_center(poly, origin, spacing_m, include_border)

# -----------------------
# Gráfica rápida
# -----------------------
xp, yp = poly.exterior.xy
plt.figure(figsize=(6,5))
plt.fill(xp, yp, alpha=0.4, fc='lightblue', ec='darkblue', label='Polígono')
if puntos:
    arr = np.array(puntos)
    plt.scatter(arr[:, 0], arr[:, 1], c='green', s=15, label='Puntos dentro')
plt.scatter(origin[0], origin[1], c='red', s=50, label='Punto inicial', zorder=5)
plt.legend()
plt.title('Retículo hexagonal optimizado, d = ' + str(spacing_m) + ' m')
plt.xlabel('X')
plt.ylabel('Y')
plt.axis('equal')
plt.grid(True)
plt.show()




# -----------------------
# Exportaciones CSV y DXF
# -----------------------
points_filename = 'puntos_hex_' + str(int(spacing_m)) + 'm.csv'
poly_filename = 'poligono_vertices.csv'
dxf_filename = 'reticulo_poligono_' + str(int(spacing_m)) + 'm.dxf'

pd.DataFrame(puntos, columns=['X', 'Y']).to_csv(points_filename, index=False)
pd.DataFrame(vertices, columns=['X', 'Y']).to_csv(poly_filename, index=False)

doc = ezdxf.new(setup=True)
ms = doc.modelspace()
poly_coords = list(vertices)
if poly_coords[0] != poly_coords[-1]:
    poly_coords.append(poly_coords[0])
ms.add_lwpolyline(poly_coords, format='xy', dxfattribs={'closed': True})
for x, y in puntos:
    ms.add_point((x, y))
doc.saveas(dxf_filename)

print('Exportado CSV puntos:', points_filename)
print('Exportado CSV polígono:', poly_filename)
print('Exportado DXF:', dxf_filename)
print('Total puntos (incluye el inicial si cae dentro):', len(puntos))

# -----------------------
# Informe Word
# -----------------------
# Guardar imágenes para el reporte
fig1_path = 'fig_poligono_origen.png'
plt.figure(figsize=(6,5))
plt.fill(xp, yp, alpha=0.4, fc='lightblue', ec='darkblue')
plt.scatter(origin[0], origin[1], c='red', s=50, label='Punto inicial')
plt.title('Polígono y punto inicial')
plt.xlabel('X')
plt.ylabel('Y')
plt.axis('equal')
plt.grid(True)
plt.legend()
plt.savefig(fig1_path, dpi=180, bbox_inches='tight')
plt.close()

fig2_path = 'fig_poligono_puntos.png'
plt.figure(figsize=(6,5))
plt.fill(xp, yp, alpha=0.4, fc='lightblue', ec='darkblue')
if puntos:
    arr = np.array(puntos)
    plt.scatter(arr[:,0], arr[:,1], c='green', s=15, label='Puntos dentro')
plt.scatter(origin[0], origin[1], c='red', s=50, label='Punto inicial')
plt.title('Puntos dentro del polígono (d = ' + str(int(spacing_m)) + ' m)')
plt.xlabel('X')
plt.ylabel('Y')
plt.axis('equal')
plt.grid(True)
plt.legend()
plt.savefig(fig2_path, dpi=180, bbox_inches='tight')
plt.close()

report_name = 'informe_reticulo_hex_' + str(int(spacing_m)) + 'm.docx'
docx = Document()
docx.add_heading('Informe de retículo hexagonal en polígono', level=1)

docx.add_heading('Datos de entrada', level=2)
docx.add_paragraph('Archivo: ' + excel_filename).alignment = WD_ALIGN_PARAGRAPH.LEFT
docx.add_paragraph('Vértices del polígono (X, Y):')
vertex_table = docx.add_table(rows=1, cols=2)
hdr = vertex_table.rows[0].cells
hdr[0].text = 'X'
hdr[1].text = 'Y'
for x, y in vertices:
    r = vertex_table.add_row().cells
    r[0].text = str(x)
    r[1].text = str(y)
docx.add_paragraph('Punto inicial del retículo (X, Y): ' + str(round(origin[0], 3)) + ', ' + str(round(origin[1], 3)))
docx.add_paragraph('Polígono y punto inicial:')
docx.add_picture(fig1_path, width=Inches(5.5))

docx.add_heading('Resultados', level=2)
docx.add_paragraph('Separación del retículo: ' + str(int(spacing_m)) + ' m')
docx.add_paragraph('Número de puntos dentro del polígono (incluyendo el inicial): ' + str(len(puntos)))
docx.add_paragraph('Polígono con los puntos dentro (punto inicial en rojo):')
docx.add_picture(fig2_path, width=Inches(5.5))

docx.save(report_name)
print('Informe Word:', report_name)