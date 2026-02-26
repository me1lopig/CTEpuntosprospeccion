import streamlit as st
import numpy as np
import matplotlib.pyplot as plt
import math
import utm
import pandas as pd
from shapely.geometry import Polygon as ShapelyPolygon, Point as ShapelyPoint
import contextily as ctx
import io
import tempfile
import ezdxf
from docx import Document
from docx.shared import Inches
import folium
from folium.plugins import Draw
from streamlit_folium import st_folium

st.set_page_config(page_title="Suite Topográfica PRO", layout="wide")
st.title("📍 Suite Topográfica: Replanteo y Optimización")

# ==========================================
# --- FUNCIONES DE MEMORIA Y ESTADO ---
# ==========================================

def limpiar_descargas():
    if 'archivos_listos' in st.session_state:
        del st.session_state['archivos_listos']

def mover_malla(dx, dy):
    st.session_state['off_x'] += dx
    st.session_state['off_y'] += dy
    limpiar_descargas()

def resetear_malla():
    st.session_state['off_x'] = 0.0
    st.session_state['off_y'] = 0.0
    limpiar_descargas()

# ESTADO ROBUSTO DESACOPLADO
if 'map_center_internal' not in st.session_state: st.session_state['map_center_internal'] = [39.5, -3.0] 
if 'map_zoom_internal' not in st.session_state: st.session_state['map_zoom_internal'] = 6
if 'txt_x' not in st.session_state: st.session_state['txt_x'] = 450000.0
if 'txt_y' not in st.session_state: st.session_state['txt_y'] = 4370000.0
if 'txt_h' not in st.session_state: st.session_state['txt_h'] = 30
if 'off_x' not in st.session_state: st.session_state['off_x'] = 0.0
if 'off_y' not in st.session_state: st.session_state['off_y'] = 0.0

if 'rebuild_map' not in st.session_state: st.session_state['rebuild_map'] = True
if 'map_obj' not in st.session_state: st.session_state['map_obj'] = None
if 'old_capa' not in st.session_state: st.session_state['old_capa'] = "Satélite PNOA"

def centrar_mapa_desde_texto():
    try:
        lat, lon = utm.to_latlon(st.session_state['txt_x'], st.session_state['txt_y'], st.session_state['txt_h'], northern=True)
        st.session_state['map_center_internal'] = [lat, lon]
        st.session_state['rebuild_map'] = True
    except:
        pass

# ==========================================
# --- EXPORTACIÓN MULTIFORMATO ---
# ==========================================

def generar_excel(df_poligono, df_puntos):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        data_poly = []
        for i, p in enumerate(df_poligono):
            lon, lat = p[0], p[1]
            e, n, zn, zl = utm.from_latlon(lat, lon)
            data_poly.append({
                'Vértice': i+1, 
                'Latitud': lat, 
                'Longitud': lon,
                'UTM_X': round(e, 3),
                'UTM_Y': round(n, 3),
                'Huso': f"{zn}{zl}"
            })
        df_poly_export = pd.DataFrame(data_poly)
        df_poly_export.to_excel(writer, sheet_name='Vertices_Parcela', index=False)
        df_puntos.to_excel(writer, sheet_name='Puntos_Replanteo', index=False)
    return output.getvalue()

def generar_informe_word(area_ha, area_m2, num_puntos, pts_ha, dist, marg, metodo, angulo_opt, off_x, off_y, fig):
    doc = Document()
    doc.add_heading('INFORME TÉCNICO DE REPLANTEO', 0)
    
    doc.add_heading('1. Datos de la Parcela:', level=1)
    doc.add_paragraph(f"• Superficie total: {area_ha:.4f} ha ({area_m2:,.2f} m²)")
    doc.add_paragraph(f"• Puntos a replantear: {num_puntos} (Densidad: {pts_ha:.0f} pts/ha)")

    doc.add_heading('2. Configuración de Malla:', level=1)
    doc.add_paragraph(f"• Método: {metodo}")
    doc.add_paragraph(f"• Separación: {dist:.2f} m | Margen al linde: {marg:.2f} m")
    if "OPTIMIZADO" in metodo:
        doc.add_paragraph(f"• Ángulo de rotación óptimo calculado: {angulo_opt}º")
    doc.add_paragraph(f"• Desplazamiento manual de ajuste: X={off_x:+.2f}m, Y={off_y:+.2f}m")

    doc.add_heading('3. Plano de Distribución:', level=1)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
        fig.savefig(tmp.name, format="png", bbox_inches="tight", dpi=300)
        doc.add_picture(tmp.name, width=Inches(6.0))
        
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

def generar_dxf(df_poligono, df_puntos):
    doc = ezdxf.new('R2010')
    
    # Configurar visualización global de puntos (Cruz 'X' visible)
    doc.header['$PDMODE'] = 3
    doc.header['$PDSIZE'] = 1.0 
    
    # Crear estructura de capas profesionales
    doc.layers.add(name="01_PARCELA_LINDE", color=3) # Verde
    doc.layers.add(name="02_PUNTOS_REPLANTEO", color=1) # Rojo
    doc.layers.add(name="03_ETIQUETAS_ID", color=2) # Amarillo
    
    msp = doc.modelspace()
    
    # 1. Dibujar el Polígono
    poly_utm = []
    for p in df_poligono:
        e, n, _, _ = utm.from_latlon(p[1], p[0])
        poly_utm.append((e, n))
        
    if poly_utm:
        msp.add_lwpolyline(poly_utm, close=True, dxfattribs={'layer': '01_PARCELA_LINDE'})
        
    # 2. Dibujar Puntos y Etiquetas
    for _, row in df_puntos.iterrows():
        x, y = row['UTM_X'], row['UTM_Y']
        id_pt = str(row['ID'])
        
        msp.add_point((x, y), dxfattribs={'layer': '02_PUNTOS_REPLANTEO'})
        msp.add_text(id_pt, dxfattribs={'layer': '03_ETIQUETAS_ID', 'height': 0.8}).set_placement((x + 0.5, y + 0.5))
        
    with tempfile.NamedTemporaryFile(suffix=".dxf", delete=False) as tmp:
        doc.saveas(tmp.name)
    with open(tmp.name, "rb") as f:
        data = f.read()
    return data

# ==========================================
# --- MOTOR DE CÁLCULO ---
# ==========================================

def procesar_poligono(coords_geo, dist, marg, mapa_fondo, opacidad, metodo, off_x, off_y):
    lat_media = np.mean(coords_geo[:, 1])
    lon_media = np.mean(coords_geo[:, 0])
    e_c, n_c, h_n, h_l = utm.from_latlon(lat_media, lon_media)
    
    lat_rad = np.radians(lat_media)
    m_deg_lat = 111132.92 - 559.82 * np.cos(2 * lat_rad)
    m_deg_lon = 111412.84 * np.cos(lat_rad) - 93.5 * np.cos(3 * lat_rad)
    
    min_lon, min_lat = np.min(coords_geo, axis=0)
    max_lon, max_lat = np.max(coords_geo, axis=0)
    
    poly_m = np.column_stack(((coords_geo[:, 0] - min_lon) * m_deg_lon, (coords_geo[:, 1] - min_lat) * m_deg_lat))
    poligono_base = ShapelyPolygon(poly_m)
    area_m2 = poligono_base.area

    poligono_util = poligono_base.buffer(-marg) if marg > 0 else poligono_base
    if poligono_util.is_empty: 
        return None, None, "El margen de seguridad es tan grande que elimina la superficie de la parcela.", 0, 0

    cx, cy = np.mean(poly_m[:, 0]), np.mean(poly_m[:, 1])
    R = math.hypot((max_lon - min_lon)*m_deg_lon, (max_lat - min_lat)*m_deg_lat)
    
    dy = dist * math.sin(math.pi/3)
    dx = dist
    
    pasos_y = int((2 * R) / dy) + 4
    pasos_x = int((2 * R) / dx) + 4
    start_x, start_y = cx - R, cy - R
    
    puntos_base = []
    for fila in range(pasos_y):
        y = start_y + fila * dy
        offset_x_tri = (dist / 2) if fila % 2 == 1 else 0
        for col in range(pasos_x):
            x = start_x + col * dx + offset_x_tri
            puntos_base.append([x, y])
    puntos_base = np.array(puntos_base)

    mejor_angulo = 0
    mejores_puntos_finales = []
    max_puntos_dentro = -1

    angulos_a_probar = range(0, 60, 1) if "OPTIMIZADO" in metodo else [0]
    
    for angulo in angulos_a_probar:
        angulo_rad = math.radians(angulo)
        cos_a, sin_a = math.cos(angulo_rad), math.sin(angulo_rad)
        
        nx = cx + (puntos_base[:, 0] - cx) * cos_a - (puntos_base[:, 1] - cy) * sin_a + off_x
        ny = cy + (puntos_base[:, 0] - cx) * sin_a + (puntos_base[:, 1] - cy) * cos_a + off_y
        
        puntos_rotados = np.column_stack((nx, ny))
        mascara = np.array([poligono_util.contains(ShapelyPoint(pt[0], pt[1])) for pt in puntos_rotados])
        puntos_validos = puntos_rotados[mascara]
        
        if len(puntos_validos) > max_puntos_dentro:
            max_puntos_dentro = len(puntos_validos)
            mejores_puntos_finales = puntos_validos
            mejor_angulo = angulo

    # PROTOCOLO DE RESCATE (Para parcelas diminutas o márgenes muy grandes)
    if len(mejores_puntos_finales) == 0:
        centro = poligono_util.representative_point()
        puntos_emergencia = [[centro.x, centro.y]]
        
        coords_ext = list(poligono_util.exterior.coords)
        distancias = [(math.hypot(px - centro.x, py - centro.y), px, py) for px, py in coords_ext]
        distancias.sort(reverse=True, key=lambda x: x[0])
        
        if len(distancias) > 0:
            p1_x, p1_y = distancias[0][1], distancias[0][2]
            puntos_emergencia.append([(centro.x + p1_x)/2, (centro.y + p1_y)/2])
            
            if len(distancias) > 1:
                p2_x, p2_y = distancias[1][1], distancias[1][2]
                v1_x, v1_y = p1_x - centro.x, p1_y - centro.y
                for d, px, py in distancias[1:]:
                    vx, vy = px - centro.x, py - centro.y
                    if (v1_x * vx + v1_y * vy) < 0:
                        p2_x, p2_y = px, py
                        break
                puntos_emergencia.append([(centro.x + p2_x)/2, (centro.y + p2_y)/2])
                
        mejores_puntos_finales = np.array(puntos_emergencia)
        metodo = "RESCATE (3 Puntos)"

    p_lon_final = (mejores_puntos_finales[:, 0] / m_deg_lon) + min_lon
    p_lat_final = (mejores_puntos_finales[:, 1] / m_deg_lat) + min_lat
    
    data_pts = []
    for lo, la in zip(p_lon_final, p_lat_final):
        e, n, zn, zl = utm.from_latlon(la, lo)
        data_pts.append({'ID': len(data_pts)+1, 'Latitud': la, 'Longitud': lo, 'UTM_X': round(e, 3), 'UTM_Y': round(n, 3), 'Huso': f"{zn}{zl}"})
    df_puntos_final = pd.DataFrame(data_pts)
    
    fig, ax = plt.subplots(figsize=(5, 3.8))
    utm_poly = np.array([[utm.from_latlon(la, lo)[0], utm.from_latlon(la, lo)[1]] for lo, la in coords_geo])
    
    ax.plot(utm_poly[:,0], utm_poly[:,1], 'k-', lw=2, label="Linde Real")
    ax.fill(utm_poly[:,0], utm_poly[:,1], alpha=opacidad, color='cyan')
    ax.scatter(df_puntos_final['UTM_X'], df_puntos_final['UTM_Y'], c='red', s=12, edgecolor='black', lw=0.4, zorder=5)
    
    if "RESCATE" in metodo:
        titulo_mapa = "⚠️ Parcela pequeña: Modo Rescate (3 Puntos)"
    else:
        titulo_mapa = f"Replanteo: {dist}m | Ajuste: X={off_x:+.2f}m, Y={off_y:+.2f}m"
        if "OPTIMIZADO" in metodo:
            titulo_mapa += f"\n(Optimizado: Rotado {mejor_angulo}º)"
    
    ax.set_title(titulo_mapa, pad=10, fontsize=9)
    ax.set_aspect('equal')
    
    xmin, xmax = ax.get_xlim()
    ymin, ymax = ax.get_ylim()
    w = xmax - xmin
    h = ymax - ymin
    
    if mapa_fondo != "Ninguno":
        # Alejamos la cámara virtualmente
        ax.set_xlim(xmin - (w * 0.5), xmax + (w * 0.5))
        ax.set_ylim(ymin - (h * 0.5), ymax + (h * 0.5))
        
        epsg_code = 32600 + h_n
        fuente = "https://www.ign.es/wmts/pnoa-ma?request=GetTile&service=WMTS&version=1.0.0&Layer=OI.OrthoimageCoverage&Style=default&Format=image/jpeg&TileMatrixSet=GoogleMapsCompatible&TileMatrix={z}&TileRow={y}&TileCol={x}" if mapa_fondo == "Satélite PNOA" else "https://www.ign.es/wmts/mapa-raster?request=getTile&layer=MTN&TileMatrixSet=GoogleMapsCompatible&TileMatrix={z}&TileCol={x}&TileRow={y}&format=image/jpeg"
        
        zoom_max = 19 if mapa_fondo == "Satélite PNOA" else 18
        
        try: 
            ctx.add_basemap(ax, crs=f"EPSG:{epsg_code}", source=fuente, alpha=1.0, reset_extent=False)
        except: 
            try:
                ctx.add_basemap(ax, crs=f"EPSG:{epsg_code}", source=fuente, alpha=1.0, reset_extent=False, zoom=zoom_max)
            except:
                pass
        
    ax.set_xlim(xmin - (w * 0.15), xmax + (w * 0.15))
    ax.set_ylim(ymin - (h * 0.15), ymax + (h * 0.15))
    
    # --- LA SOLUCIÓN A LOS NÚMEROS SOLAPADOS EN EL EJE X ---
    ax.tick_params(axis='y', which='major', labelsize=7)
    # Le añadimos labelrotation=90 al eje x para que se lean de abajo hacia arriba
    ax.tick_params(axis='x', which='major', labelsize=7, labelrotation=90)
    ax.ticklabel_format(useOffset=False, style='plain')
        
    return df_puntos_final, fig, None, area_m2, mejor_angulo, ("RESCATE" in metodo)

# ==========================================
# --- INTERFAZ (TABS SEPARADAS) ---
# ==========================================

tab1, tab2 = st.tabs(["🗺️ 1. DIBUJO Y LOCALIZACIÓN", "⚙️ 2. CONFIGURACIÓN Y RESULTADOS"])

# ------------------------------------------
# PESTAÑA 1
# ------------------------------------------
with tab1:
    col_ctrl, col_map = st.columns([1, 4])
    
    with col_map:
        if st.session_state['rebuild_map'] or st.session_state['map_obj'] is None:
            m = folium.Map(location=st.session_state['map_center_internal'], zoom_start=st.session_state['map_zoom_internal'], max_zoom=24)
            
            url_pnoa = "https://www.ign.es/wmts/pnoa-ma?request=GetTile&service=WMTS&version=1.0.0&Layer=OI.OrthoimageCoverage&Style=default&Format=image/jpeg&TileMatrixSet=GoogleMapsCompatible&TileMatrix={z}&TileRow={y}&TileCol={x}"
            url_mtn = "https://www.ign.es/wmts/mapa-raster?request=getTile&layer=MTN&TileMatrixSet=GoogleMapsCompatible&TileMatrix={z}&TileCol={x}&TileRow={y}&format=image/jpeg"
            
            capa_seleccionada = st.session_state['old_capa']
            zoom_nativo_maximo = 19 if capa_seleccionada == "Satélite PNOA" else 18
            
            folium.TileLayer(
                tiles=(url_pnoa if capa_seleccionada == "Satélite PNOA" else url_mtn), 
                attr="IGN",
                max_native_zoom=zoom_nativo_maximo,
                max_zoom=24
            ).add_to(m)

            if 'poligono_usuario' in st.session_state:
                coords_f = [[p[1], p[0]] for p in st.session_state['poligono_usuario']]
                folium.Polygon(locations=coords_f, color="#FFD700", fill=True, fill_opacity=0.3).add_to(m)

            Draw(export=False, position='topleft', draw_options={'polyline':False, 'rectangle':False, 'circle':False, 'marker':False, 'circlemarker':False}).add_to(m)
            
            st.session_state['map_obj'] = m
            st.session_state['rebuild_map'] = False

        output_mapa = st_folium(st.session_state['map_obj'], width="100%", height=500, key="visor_principal", returned_objects=["center", "zoom", "all_drawings"])

        if output_mapa and output_mapa.get("center"):
            c_lat, c_lon = output_mapa["center"]["lat"], output_mapa["center"]["lng"]
            c_zoom = output_mapa.get("zoom", st.session_state['map_zoom_internal'])
            
            old_lat, old_lon = st.session_state['map_center_internal']
            
            if abs(c_lat - old_lat) > 0.0001 or abs(c_lon - old_lon) > 0.0001 or c_zoom != st.session_state['map_zoom_internal']:
                st.session_state['map_center_internal'] = [c_lat, c_lon]
                st.session_state['map_zoom_internal'] = c_zoom
                
                try:
                    e_new, n_new, h_new, _ = utm.from_latlon(c_lat, c_lon)
                    st.session_state['txt_x'] = round(e_new, 2)
                    st.session_state['txt_y'] = round(n_new, 2)
                    st.session_state['txt_h'] = h_new
                except: pass

        if output_mapa and output_mapa.get("all_drawings") and len(output_mapa["all_drawings"]) > 0:
            nuevas_coords = output_mapa["all_drawings"][-1]["geometry"]["coordinates"][0]
            if 'poligono_usuario' not in st.session_state or not np.array_equal(st.session_state['poligono_usuario'], np.array(nuevas_coords)):
                st.session_state['poligono_usuario'] = np.array(nuevas_coords)
                st.toast("✅ Polígono guardado. Ve a la Pestaña 2.")
                st.session_state['rebuild_map'] = True
                st.rerun()

    with col_ctrl:
        st.subheader("Centro del Mapa (UTM)")
        
        ux_in = st.number_input("UTM Este (X):", value=float(st.session_state['txt_x']), format="%.2f", step=100.0)
        uy_in = st.number_input("UTM Norte (Y):", value=float(st.session_state['txt_y']), format="%.2f", step=100.0)
        h_in = st.number_input("Huso:", min_value=28, max_value=31, value=int(st.session_state['txt_h']), step=1)
        
        if ux_in != st.session_state['txt_x'] or uy_in != st.session_state['txt_y'] or h_in != st.session_state['txt_h']:
            st.session_state['txt_x'] = ux_in
            st.session_state['txt_y'] = uy_in
            st.session_state['txt_h'] = h_in
            try:
                lat, lon = utm.to_latlon(ux_in, uy_in, h_in, northern=True)
                st.session_state['map_center_internal'] = [lat, lon]
                st.session_state['rebuild_map'] = True
                st.rerun()
            except: pass
        
        st.divider()
        capa_elegida = st.radio("Capa Base IGN:", ["Satélite PNOA", "Topográfico MTN"], key="capa_base_selector")
        
        if capa_elegida != st.session_state['old_capa']:
            st.session_state['old_capa'] = capa_elegida
            st.session_state['rebuild_map'] = True
            st.rerun()

        if st.button("🗑️ Borrar Polígono", use_container_width=True):
            if 'poligono_usuario' in st.session_state:
                del st.session_state['poligono_usuario']
            st.session_state['rebuild_map'] = True
            st.rerun()

# ------------------------------------------
# PESTAÑA 2
# ------------------------------------------
with tab2:
    if 'poligono_usuario' not in st.session_state:
        st.warning("⚠️ Ve a la Pestaña 1 y dibuja un polígono para poder configurar la malla.")
    else:
        col_conf, col_res = st.columns([1.2, 2.8])
        
        with col_conf:
            st.header("⚙️ Geometría")
            metodo_dist = st.selectbox("Método de Distribución:", ["Hexagonal Normal (Norte-Sur)", "Hexagonal OPTIMIZADO (Búsqueda del Máximo)"], on_change=limpiar_descargas)
            distancia = st.number_input("📏 Distancia (m):", min_value=1.0, value=25.0, step=0.5, on_change=limpiar_descargas)
            margen = st.number_input("🛡️ Margen (m):", min_value=0.0, value=1.0, step=0.5, on_change=limpiar_descargas)
            
            st.divider()
            st.header("🎮 Ajuste Fino")
            paso = st.selectbox("Resolución del botón:", [1.0, 0.5, 0.1, 0.01], format_func=lambda x: f"{x} m", help="Movimiento de malla por clic.")
            
            st.markdown("<div style='text-align: center; margin-bottom: 5px; font-size: 0.9em;'>Desplazamiento manual: <br><b>X: {:.2f}m | Y: {:.2f}m</b></div>".format(st.session_state['off_x'], st.session_state['off_y']), unsafe_allow_html=True)
            
            c1, c2, c3 = st.columns([1, 1.2, 1])
            with c2: st.button("⬆️ N", on_click=mover_malla, args=(0, paso), use_container_width=True)
            
            c4, c5, c6 = st.columns([1, 1.2, 1])
            with c4: st.button("⬅️ O", on_click=mover_malla, args=(-paso, 0), use_container_width=True)
            with c5: st.button("🔄", on_click=resetear_malla, use_container_width=True, help="Resetear a 0")
            with c6: st.button("➡️ E", on_click=mover_malla, args=(paso, 0), use_container_width=True)
            
            c7, c8, c9 = st.columns([1, 1.2, 1])
            with c8: st.button("⬇️ S", on_click=mover_malla, args=(0, -paso), use_container_width=True)

            st.divider()
            mapa_final = st.radio("Fondo Plano Final:", ["Satélite PNOA", "Topográfico MTN", "Ninguno"], on_change=limpiar_descargas)
            opacidad_final = st.slider("Opacidad Parcela:", 0.0, 1.0, 0.3)

        with col_res:
            texto_carga = 'Calculando rotación óptima...' if 'OPTIMIZADO' in metodo_dist else 'Procesando Malla...'
            with st.spinner(texto_carga):
                df_res, fig_final, error, area_m2, angulo_opt, is_rescate = procesar_poligono(
                    st.session_state['poligono_usuario'], distancia, margen, 
                    mapa_final, opacidad_final, metodo_dist, st.session_state['off_x'], st.session_state['off_y']
                )

            if error:
                st.error(error)
            else:
                if is_rescate:
                    st.warning("⚠️ El área de la parcela es muy pequeña para la distancia requerida. Se han generado automáticamente 3 puntos de control.")
                
                m1, m2, m3 = st.columns(3)
                if "OPTIMIZADO" in metodo_dist and not is_rescate:
                    m1.metric("Puntos", f"{len(df_res)} pts", delta=f"Rotado {angulo_opt}º", delta_color="normal")
                else:
                    m1.metric("Puntos", f"{len(df_res)} pts")
                m2.metric("Área Útil", f"{area_m2/10000:.2f} ha")
                m3.metric("Densidad", f"{len(df_res)/(area_m2/10000):.0f} pts/ha")
                
                st.pyplot(fig_final, use_container_width=False)
                
                if not st.session_state.get('archivos_listos'):
                    if st.button("🚀 PREPARAR INFORMES (Excel, Word y DXF)", type="primary"):
                        st.session_state['excel_data'] = generar_excel(st.session_state['poligono_usuario'], df_res)
                        st.session_state['word_data'] = generar_informe_word(area_m2/10000, area_m2, len(df_res), len(df_res)/(area_m2/10000), distancia, margen, metodo_dist, angulo_opt, st.session_state['off_x'], st.session_state['off_y'], fig_final)
                        st.session_state['dxf_data'] = generar_dxf(st.session_state['poligono_usuario'], df_res)
                        st.session_state['archivos_listos'] = True
                        st.rerun()
                
                if st.session_state.get('archivos_listos'):
                    cb1, cb2, cb3 = st.columns(3)
                    cb1.download_button("📊 Descargar Excel", st.session_state['excel_data'], "Coordenadas_Replanteo.xlsx", use_container_width=True)
                    cb2.download_button("📝 Descargar Word", st.session_state['word_data'], "Informe_Topografico.docx", use_container_width=True)
                    cb3.download_button("📐 Descargar DXF", st.session_state['dxf_data'], "Plano_CAD_Replanteo.dxf", use_container_width=True)