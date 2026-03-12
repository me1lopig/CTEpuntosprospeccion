import streamlit as st
import numpy as np
import matplotlib.pyplot as plt
import math
import utm
import pandas as pd
from shapely.geometry import Polygon as ShapelyPolygon, Point as ShapelyPoint
import contextily as ctx
import io
import tempfile
import ezdxf
from docx import Document
from docx.shared import Inches
import folium
from folium.plugins import Draw
from streamlit_folium import st_folium

st.set_page_config(page_title="Distribución de trabajos de Campo a realizar", layout="wide")
st.title("📍 Distribución de trabajos de campo")

# ==========================================
# --- FUNCIONES DE MEMORIA Y ESTADO ---
# ==========================================

def limpiar_archivos():
    if 'archivos_listos' in st.session_state:
        del st.session_state['archivos_listos']

def resetear_calculo():
    if 'df_puntos_actual' in st.session_state:
        del st.session_state['df_puntos_actual']
    limpiar_archivos()

def mover_malla(dx, dy):
    st.session_state['off_x'] += dx
    st.session_state['off_y'] += dy
    resetear_calculo()

def resetear_malla():
    st.session_state['off_x'] = 0.0
    st.session_state['off_y'] = 0.0
    resetear_calculo()

def alternar_anclaje():
    st.session_state['rebuild_map'] = True

# ESTADO ROBUSTO DESACOPLADO
if 'map_center_internal' not in st.session_state: st.session_state['map_center_internal'] = [39.5, -3.0] 
if 'map_zoom_internal' not in st.session_state: st.session_state['map_zoom_internal'] = 6
if 'txt_x' not in st.session_state: st.session_state['txt_x'] = 450000.0
if 'txt_y' not in st.session_state: st.session_state['txt_y'] = 4370000.0
if 'txt_h' not in st.session_state: st.session_state['txt_h'] = 30

# Nuevas variables en memoria para sincronizar casillas y mapa
if 'in_utm_x' not in st.session_state: st.session_state['in_utm_x'] = 450000.0
if 'in_utm_y' not in st.session_state: st.session_state['in_utm_y'] = 4370000.0
if 'in_utm_h' not in st.session_state: st.session_state['in_utm_h'] = 30

if 'off_x' not in st.session_state: st.session_state['off_x'] = 0.0
if 'off_y' not in st.session_state: st.session_state['off_y'] = 0.0

if 'rebuild_map' not in st.session_state: st.session_state['rebuild_map'] = True
if 'map_obj' not in st.session_state: st.session_state['map_obj'] = None
if 'old_capa' not in st.session_state: st.session_state['old_capa'] = "Satélite PNOA"
if 'mapa_anclado' not in st.session_state: st.session_state['mapa_anclado'] = False

def centrar_mapa_desde_texto():
    try:
        lat, lon = utm.to_latlon(st.session_state['txt_x'], st.session_state['txt_y'], st.session_state['txt_h'], northern=True)
        st.session_state['map_center_internal'] = [lat, lon]
        st.session_state['rebuild_map'] = True
    except:
        pass

# ==========================================
# --- EXPORTACIÓN MULTIFORMATO ---
# ==========================================

def generar_excel(df_poligono, df_puntos):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        data_poly = []
        for i, p in enumerate(df_poligono):
            lon, lat = p[0], p[1]
            e, n, zn, zl = utm.from_latlon(lat, lon)
            data_poly.append({
                'Vértice': i+1, 
                'Latitud': lat, 
                'Longitud': lon,
                'UTM_X': round(e, 3),
                'UTM_Y': round(n, 3),
                'Huso': f"{zn}{zl}"
            })
        df_poly_export = pd.DataFrame(data_poly)
        df_poly_export.to_excel(writer, sheet_name='Vertices_Parcela', index=False)
        
        # Como df_puntos ya tiene Latitud y Longitud calculados, solo ordenamos las columnas
        df_puntos_export = df_puntos.copy()
        cols = ['ID', 'Latitud', 'Longitud', 'UTM_X', 'UTM_Y', 'Huso']
        df_puntos_export = df_puntos_export[cols]
        df_puntos_export.to_excel(writer, sheet_name='Puntos_Replanteo', index=False)
        
    return output.getvalue()

def generar_informe_word(area_ha, area_m2, num_puntos, pts_ha, dist, marg, metodo, angulo_opt, off_x, off_y, fig):
    doc = Document()
    doc.add_heading('INFORME TÉCNICO DE REPLANTEO', 0)
    
    doc.add_heading('1. Datos de la Parcela:', level=1)
    doc.add_paragraph(f"• Superficie total: {area_ha:.4f} ha ({area_m2:,.2f} m²)")
    doc.add_paragraph(f"• Puntos a replantear: {num_puntos} (Densidad: {pts_ha:.0f} pts/ha)")

    doc.add_heading('2. Configuración de Malla:', level=1)
    doc.add_paragraph(f"• Método: {metodo}")
    doc.add_paragraph(f"• Separación: {dist:.2f} m | Distancia al borde {marg:.2f} m")
    if "OPTIMIZADO" in metodo:
        doc.add_paragraph(f"• Ángulo de rotación óptimo calculado: {angulo_opt}º")
    doc.add_paragraph(f"• Desplazamiento manual de ajuste: X={off_x:+.2f}m, Y={off_y:+.2f}m")

    doc.add_heading('3. Plano de Distribución:', level=1)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
        fig.savefig(tmp.name, format="png", bbox_inches="tight", dpi=300)
        doc.add_picture(tmp.name, width=Inches(6.0))
        
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

def generar_dxf(df_poligono, df_puntos, h_n):
    doc = ezdxf.new('R2010')
    doc.header['$PDMODE'] = 3
    doc.header['$PDSIZE'] = 1.0 
    
    doc.layers.add(name="01_PARCELA_BORDE", color=3) 
    doc.layers.add(name="02_PUNTOS_REPLANTEO", color=1) 
    doc.layers.add(name="03_ETIQUETAS_ID", color=2) 
    
    msp = doc.modelspace()
    
    # FORZAMOS LA ZONA DEL DXF PARA EVITAR SALTOS EN AUTOCAD
    poly_utm = []
    for p in df_poligono:
        e, n, _, _ = utm.from_latlon(p[1], p[0], force_zone_number=h_n)
        poly_utm.append((e, n))
        
    if poly_utm:
        msp.add_lwpolyline(poly_utm, close=True, dxfattribs={'layer': '01_PARCELA_BORDE'})
        
    for _, row in df_puntos.iterrows():
        # Recuperamos la posición forzada para que cuadre exactamente dentro del polígono en CAD
        x, y, _, _ = utm.from_latlon(row['Latitud'], row['Longitud'], force_zone_number=h_n)
        id_pt = str(row['ID'])
        
        msp.add_point((x, y), dxfattribs={'layer': '02_PUNTOS_REPLANTEO'})
        msp.add_text(id_pt, dxfattribs={'layer': '03_ETIQUETAS_ID', 'height': 0.8}).set_placement((x + 0.5, y + 0.5))
        
    with tempfile.NamedTemporaryFile(suffix=".dxf", delete=False) as tmp:
        doc.saveas(tmp.name)
    with open(tmp.name, "rb") as f:
        data = f.read()
    return data

# ==========================================
# --- MOTOR DE CÁLCULO MEJORADO (PROYECCIÓN FORZADA) ---
# ==========================================

def calcular_malla(coords_geo, dist, marg, metodo, off_x, off_y):
    lat_media = np.mean(coords_geo[:, 1])
    lon_media = np.mean(coords_geo[:, 0])
    
    # 1. Determinamos el huso central ("Base") de la parcela
    _, _, h_n, h_l = utm.from_latlon(lat_media, lon_media)
    
    # 2. PROYECCIÓN FORZADA: Convertimos todo el polígono al Huso Central exacto en metros
    poly_m = np.array([utm.from_latlon(pt[1], pt[0], force_zone_number=h_n)[:2] for pt in coords_geo])
    
    poligono_base = ShapelyPolygon(poly_m)
    area_m2 = poligono_base.area

    poligono_util = poligono_base.buffer(-marg) if marg > 0 else poligono_base
    if poligono_util.is_empty: 
        return None, "El margen de seguridad es tan grande que elimina la superficie de la parcela.", 0, 0, False, h_n

    cx, cy = np.mean(poly_m[:, 0]), np.mean(poly_m[:, 1])
    min_x, min_y = np.min(poly_m, axis=0)
    max_x, max_y = np.max(poly_m, axis=0)
    R = math.hypot(max_x - min_x, max_y - min_y)
    
    dy = dist * math.sin(math.pi/3)
    dx = dist
    
    pasos_y = int((2 * R) / dy) + 4
    pasos_x = int((2 * R) / dx) + 4
    start_x, start_y = cx - R, cy - R
    
    puntos_base = []
    for fila in range(pasos_y):
        y = start_y + fila * dy
        offset_x_tri = (dist / 2) if fila % 2 == 1 else 0
        for col in range(pasos_x):
            x = start_x + col * dx + offset_x_tri
            puntos_base.append([x, y])
    puntos_base = np.array(puntos_base)

    mejor_angulo = 0
    mejores_puntos_finales = []
    max_puntos_dentro = -1

    angulos_a_probar = range(0, 60, 1) if "OPTIMIZADO" in metodo else [0]
    
    for angulo in angulos_a_probar:
        angulo_rad = math.radians(angulo)
        cos_a, sin_a = math.cos(angulo_rad), math.sin(angulo_rad)
        
        nx = cx + (puntos_base[:, 0] - cx) * cos_a - (puntos_base[:, 1] - cy) * sin_a + off_x
        ny = cy + (puntos_base[:, 0] - cx) * sin_a + (puntos_base[:, 1] - cy) * cos_a + off_y
        
        puntos_rotados = np.column_stack((nx, ny))
        mascara = np.array([poligono_util.contains(ShapelyPoint(pt[0], pt[1])) for pt in puntos_rotados])
        puntos_validos = puntos_rotados[mascara]
        
        if len(puntos_validos) > max_puntos_dentro:
            max_puntos_dentro = len(puntos_validos)
            mejores_puntos_finales = puntos_validos
            mejor_angulo = angulo

    is_rescate = False
    if len(mejores_puntos_finales) < 3:
        centro = poligono_util.representative_point()
        puntos_emergencia = [[centro.x, centro.y]]
        
        coords_ext = list(poligono_util.exterior.coords)
        distancias = [(math.hypot(px - centro.x, py - centro.y), px, py) for px, py in coords_ext]
        distancias.sort(reverse=True, key=lambda x: x[0])
        
        if len(distancias) > 0:
            p1_x, p1_y = distancias[0][1], distancias[0][2]
            puntos_emergencia.append([(centro.x + p1_x)/2, (centro.y + p1_y)/2])
            
            if len(distancias) > 1:
                p2_x, p2_y = distancias[1][1], distancias[1][2]
                v1_x, v1_y = p1_x - centro.x, p1_y - centro.y
                for d, px, py in distancias[1:]:
                    vx, vy = px - centro.x, py - centro.y
                    if (v1_x * vx + v1_y * vy) < 0:
                        p2_x, p2_y = px, py
                        break
                puntos_emergencia.append([(centro.x + p2_x)/2, (centro.y + p2_y)/2])
                
        mejores_puntos_finales = np.array(puntos_emergencia)
        is_rescate = True

    # 3. Extraemos los puntos y recuperamos su Lat/Lon y UTM Natural
    data_pts = []
    is_northern = (lat_media >= 0)
    for pt in mejores_puntos_finales:
        x_forz, y_forz = pt[0], pt[1]
        
        # Devolvemos el punto forzado a Lat/Lon exacto
        la, lo = utm.to_latlon(x_forz, y_forz, h_n, northern=is_northern)
        
        # Obtenemos su Huso Natural para la tabla (sin distorsión topográfica de campo)
        e_nat, n_nat, zn, zl = utm.from_latlon(la, lo)
        
        data_pts.append({
            'ID': str(len(data_pts)+1), 
            'Latitud': round(la, 6),
            'Longitud': round(lo, 6),
            'UTM_X': round(e_nat, 3), 
            'UTM_Y': round(n_nat, 3), 
            'Huso': f"{zn}{zl}"
        })
    
    df_puntos_final = pd.DataFrame(data_pts)
    return df_puntos_final, None, area_m2, mejor_angulo, is_rescate, h_n

def dibujar_plano(coords_geo, df_puntos, mapa_fondo, opacidad, metodo, dist, mejor_angulo, is_rescate, h_n, tam_letra):
    # Proyectamos forzosamente al huso central para que el dibujo de Matplotlib no se deforme
    utm_poly = np.array([utm.from_latlon(la, lo, force_zone_number=h_n)[:2] for lo, la in coords_geo])
    
    min_x, min_y = np.min(utm_poly, axis=0)
    max_x, max_y = np.max(utm_poly, axis=0)
    
    w_real = max_x - min_x
    h_real = max_y - min_y
    if w_real == 0: w_real = 1
    
    ratio = h_real / w_real
    fig_w = 6.0
    fig_h = min(max(fig_w * ratio, 4.0), 7.5) 
    
    fig, ax = plt.subplots(figsize=(fig_w, fig_h))
    
    ax.plot(utm_poly[:,0], utm_poly[:,1], 'k-', lw=2, label="Linde Real")
    ax.fill(utm_poly[:,0], utm_poly[:,1], alpha=opacidad, color='cyan')
    
    # Dibujamos los puntos recalculando temporalmente su posición forzada
    for _, row in df_puntos.iterrows():
        x_plot, y_plot, _, _ = utm.from_latlon(row['Latitud'], row['Longitud'], force_zone_number=h_n)
        ax.scatter(x_plot, y_plot, c='red', s=12, edgecolor='black', lw=0.4, zorder=5)
        ax.text(x_plot, y_plot, str(row['ID']), 
                fontsize=tam_letra, ha='center', va='center',
                bbox=dict(facecolor='white', alpha=0.8, edgecolor='black', boxstyle='round,pad=0.2', lw=0.5),
                zorder=6)
    
    if is_rescate:
        titulo_mapa = "⚠️ Parcela pequeña: (3 Puntos)"
    else:
        titulo_mapa = f"Distancia entre puntos: {dist}m "
    
    ax.set_title(titulo_mapa, pad=10, fontsize=9)
    ax.set_aspect('equal')
    
    xmin, xmax = ax.get_xlim()
    ymin, ymax = ax.get_ylim()
    w = xmax - xmin
    h = ymax - ymin
    
    if mapa_fondo != "Ninguno":
        ax.set_xlim(xmin - (w * 0.5), xmax + (w * 0.5))
        ax.set_ylim(ymin - (h * 0.5), ymax + (h * 0.5))
        
        epsg_code = 32600 + h_n
        fuente = "https://www.ign.es/wmts/pnoa-ma?request=GetTile&service=WMTS&version=1.0.0&Layer=OI.OrthoimageCoverage&Style=default&Format=image/jpeg&TileMatrixSet=GoogleMapsCompatible&TileMatrix={z}&TileRow={y}&TileCol={x}" if mapa_fondo == "Satélite PNOA" else "https://www.ign.es/wmts/mapa-raster?request=getTile&layer=MTN&TileMatrixSet=GoogleMapsCompatible&TileMatrix={z}&TileCol={x}&TileRow={y}&format=image/jpeg"
        zoom_max = 19 if mapa_fondo == "Satélite PNOA" else 18
        
        try: 
            ctx.add_basemap(ax, crs=f"EPSG:{epsg_code}", source=fuente, alpha=1.0, reset_extent=False)
        except: 
            try:
                ctx.add_basemap(ax, crs=f"EPSG:{epsg_code}", source=fuente, alpha=1.0, reset_extent=False, zoom=zoom_max)
            except:
                pass
        
    ax.set_xlim(xmin - (w * 0.15), xmax + (w * 0.15))
    ax.set_ylim(ymin - (h * 0.15), ymax + (h * 0.15))
    
    ax.tick_params(axis='y', which='major', labelsize=7)
    ax.tick_params(axis='x', which='major', labelsize=7, labelrotation=90)
    ax.ticklabel_format(useOffset=False, style='plain')
    
    fig.tight_layout()
        
    return fig

# ==========================================
# --- INTERFAZ ORIGINAL REPARADA ---
# ==========================================

tab1, tab2 = st.tabs(["🗺️ 1. DIBUJO Y LOCALIZACIÓN", "⚙️ 2. CONFIGURACIÓN Y RESULTADOS"])

# ------------------------------------------
# PESTAÑA 1
# ------------------------------------------
with tab1:
    col_ctrl, col_map = st.columns([1, 4])
    
    with col_map:
        if st.session_state['rebuild_map'] or st.session_state['map_obj'] is None:
            m = folium.Map(
                location=st.session_state['map_center_internal'], 
                zoom_start=st.session_state['map_zoom_internal'], 
                max_zoom=24,
                zoom_control=not st.session_state['mapa_anclado'],
                scrollWheelZoom=not st.session_state['mapa_anclado'],
                dragging=not st.session_state['mapa_anclado'],
                doubleClickZoom=not st.session_state['mapa_anclado']
            )
            
            url_pnoa = "https://www.ign.es/wmts/pnoa-ma?request=GetTile&service=WMTS&version=1.0.0&Layer=OI.OrthoimageCoverage&Style=default&Format=image/jpeg&TileMatrixSet=GoogleMapsCompatible&TileMatrix={z}&TileRow={y}&TileCol={x}"
            url_mtn = "https://www.ign.es/wmts/mapa-raster?request=getTile&layer=MTN&TileMatrixSet=GoogleMapsCompatible&TileMatrix={z}&TileCol={x}&TileRow={y}&format=image/jpeg"
            
            capa_seleccionada = st.session_state['old_capa']
            zoom_nativo_maximo = 19 if capa_seleccionada == "Satélite PNOA" else 18
            
            folium.TileLayer(
                tiles=(url_pnoa if capa_seleccionada == "Satélite PNOA" else url_mtn), 
                attr="IGN",
                max_native_zoom=zoom_nativo_maximo,
                max_zoom=24
            ).add_to(m)

            if 'poligono_usuario' in st.session_state:
                coords_f = [[p[1], p[0]] for p in st.session_state['poligono_usuario']]
                folium.Polygon(locations=coords_f, color="#FFD700", fill=True, fill_opacity=0.3).add_to(m)

            Draw(export=False, position='topleft', draw_options={'polyline':False, 'rectangle':False, 'circle':False, 'marker':False, 'circlemarker':False}).add_to(m)
            
            st.session_state['map_obj'] = m
            st.session_state['rebuild_map'] = False

        output_mapa = st_folium(st.session_state['map_obj'], width="100%", height=500, key="visor_principal", returned_objects=["center", "zoom", "all_drawings"])

        if output_mapa and output_mapa.get("center"):
            c_lat, c_lon = output_mapa["center"]["lat"], output_mapa["center"]["lng"]
            c_zoom = output_mapa.get("zoom", st.session_state['map_zoom_internal'])
            
            old_lat, old_lon = st.session_state['map_center_internal']
            
            if abs(c_lat - old_lat) > 0.0001 or abs(c_lon - old_lon) > 0.0001 or c_zoom != st.session_state['map_zoom_internal']:
                st.session_state['map_center_internal'] = [c_lat, c_lon]
                st.session_state['map_zoom_internal'] = c_zoom
                
                try:
                    e_new, n_new, h_new, _ = utm.from_latlon(c_lat, c_lon)
                    st.session_state['txt_x'] = round(e_new, 2)
                    st.session_state['txt_y'] = round(n_new, 2)
                    st.session_state['txt_h'] = h_new
                    
                    st.session_state['ui_utm_x'] = round(e_new, 2)
                    st.session_state['ui_utm_y'] = round(n_new, 2)
                    st.session_state['ui_utm_h'] = h_new
                except: pass

        if output_mapa and output_mapa.get("all_drawings") and len(output_mapa["all_drawings"]) > 0:
            nuevas_coords = output_mapa["all_drawings"][-1]["geometry"]["coordinates"][0]
            if 'poligono_usuario' not in st.session_state or not np.array_equal(st.session_state['poligono_usuario'], np.array(nuevas_coords)):
                st.session_state['poligono_usuario'] = np.array(nuevas_coords)
                st.toast("✅ Polígono guardado. Ve a la Pestaña 2.")
                st.session_state['rebuild_map'] = True
                resetear_calculo() 
                st.rerun()

    with col_ctrl:
        st.subheader("Centro del Mapa (UTM)")
        
        if 'ui_utm_x' not in st.session_state: st.session_state['ui_utm_x'] = st.session_state['txt_x']
        if 'ui_utm_y' not in st.session_state: st.session_state['ui_utm_y'] = st.session_state['txt_y']
        if 'ui_utm_h' not in st.session_state: st.session_state['ui_utm_h'] = st.session_state['txt_h']

        def manual_utm_change():
            st.session_state['txt_x'] = st.session_state['ui_utm_x']
            st.session_state['txt_y'] = st.session_state['ui_utm_y']
            st.session_state['txt_h'] = st.session_state['ui_utm_h']
            try:
                lat, lon = utm.to_latlon(st.session_state['ui_utm_x'], st.session_state['ui_utm_y'], st.session_state['ui_utm_h'], northern=True)
                st.session_state['map_center_internal'] = [lat, lon]
                st.session_state['rebuild_map'] = True
            except: pass

        st.number_input("UTM Este (X):", format="%.2f", step=100.0, key="ui_utm_x", on_change=manual_utm_change)
        st.number_input("UTM Norte (Y):", format="%.2f", step=100.0, key="ui_utm_y", on_change=manual_utm_change)
        st.number_input("Huso:", min_value=28, max_value=31, step=1, key="ui_utm_h", on_change=manual_utm_change)
        
        st.divider()
        
        def al_cambiar_capa():
            st.session_state['old_capa'] = st.session_state['capa_base_selector']
            st.session_state['rebuild_map'] = True
            
        st.radio("Capa Base IGN:", ["Satélite PNOA", "Topográfico MTN"], key="capa_base_selector", on_change=al_cambiar_capa)

        st.divider()
        st.toggle(
            "🔒 Anclar Mapa (Bloquea movimiento)", 
            key="mapa_anclado", 
            on_change=alternar_anclaje
        )

        st.divider()
        def borrar_poligono():
            if 'poligono_usuario' in st.session_state:
                del st.session_state['poligono_usuario']
            st.session_state['rebuild_map'] = True

        st.button("🗑️ Borrar Polígono", use_container_width=True, type="primary", on_click=borrar_poligono)

# ------------------------------------------
# PESTAÑA 2
# ------------------------------------------
with tab2:
    if 'poligono_usuario' not in st.session_state:
        st.warning("⚠️ Ve a la Pestaña 1 y dibuja un polígono para poder configurar la malla.")
    else:
        col_conf, col_res = st.columns([1, 4])
        
        with col_conf:
            st.header("⚙️ Geometría")
            metodo_dist = st.selectbox("Método:", ["Hexagonal Normal (Norte-Sur)", "Hexagonal OPTIMIZADO (Búsqueda del Máximo)"], on_change=resetear_calculo)
            distancia = st.number_input("📏 Distancia (m):", min_value=1.0, value=25.0, step=0.5, on_change=resetear_calculo)
            margen = st.number_input("🛡️ Margen (m):", min_value=0.0, value=1.0, step=0.5, on_change=resetear_calculo)
            
            st.divider()
            st.header("💎 Ajuste Fino")
            paso = st.number_input("Paso de desplazamiento (m):", min_value=0.01, value=1.00, step=0.50, format="%.2f")
            
            st.markdown("<div style='text-align: center; margin-bottom: 5px; font-size: 0.9em;'>Desplazamiento:<br><b>X: {:.2f}m | Y: {:.2f}m</b></div>".format(st.session_state['off_x'], st.session_state['off_y']), unsafe_allow_html=True)
            
            c1, c2, c3 = st.columns([1, 1.2, 1])
            with c2: st.button("⬆️ N", on_click=mover_malla, args=(0, paso), use_container_width=True)
            
            c4, c5, c6 = st.columns([1, 1.2, 1])
            with c4: st.button("⬅️ O", on_click=mover_malla, args=(-paso, 0), use_container_width=True)
            with c5: st.button("🔄", on_click=resetear_malla, use_container_width=True)
            with c6: st.button("➡️ E", on_click=mover_malla, args=(paso, 0), use_container_width=True)
            
            c7, c8, c9 = st.columns([1, 1.2, 1])
            with c8: st.button("⬇️ S", on_click=mover_malla, args=(0, -paso), use_container_width=True)

            st.divider()
            st.header("🎨 Visualización")
            mapa_final = st.radio("Fondo Final:", ["Satélite PNOA", "Topográfico MTN", "Ninguno"], on_change=limpiar_archivos)
            opacidad_final = st.slider("Opacidad Parcela:", 0.0, 1.0, 0.3, on_change=limpiar_archivos)
            tam_letra = st.slider("🔠 Tamaño ID Puntos:", 4, 20, 8, on_change=limpiar_archivos)

        with col_res:
            texto_carga = 'Calculando rotación óptima...' if 'OPTIMIZADO' in metodo_dist else 'Procesando Malla...'
            with st.spinner(texto_carga):
                if 'df_puntos_actual' not in st.session_state:
                    df_res, error, area_m2, angulo_opt, is_rescate, h_n = calcular_malla(
                        st.session_state['poligono_usuario'], distancia, margen, 
                        metodo_dist, st.session_state['off_x'], st.session_state['off_y']
                    )
                    if error:
                        st.error(error)
                        st.stop()
                    else:
                        st.session_state['df_puntos_actual'] = df_res
                        st.session_state['area_m2'] = area_m2
                        st.session_state['angulo_opt'] = angulo_opt
                        st.session_state['is_rescate'] = is_rescate
                        st.session_state['h_n'] = h_n

            df_actual = st.session_state['df_puntos_actual']
            area_m2 = st.session_state['area_m2']
            angulo_opt = st.session_state['angulo_opt']
            is_rescate = st.session_state['is_rescate']
            h_n = st.session_state['h_n']

            if is_rescate:
                st.warning("⚠️ El área de la parcela es muy pequeña o la distancia muy grande. Se han generado 3 puntos de control para orientar la base.")
            
            m1, m2, m3 = st.columns(3)
            if "OPTIMIZADO" in metodo_dist and not is_rescate:
                m1.metric("Puntos Activos", f"{len(df_actual)} pts", delta=f"Rotado {angulo_opt}º", delta_color="normal")
            else:
                m1.metric("Puntos Activos", f"{len(df_actual)} pts")
            m2.metric("Área Útil", f"{area_m2:.2f} m2")
            
            if area_m2 > 0:
                m3.metric("Densidad Resultante", f"{len(df_actual)/(area_m2/10000):.0f} pts/ha")
            else:
                m3.metric("Densidad Resultante", "0 pts/ha")
            
            col_plano, col_tabla = st.columns([1.6, 1])
            
            with col_tabla:
                st.markdown("📝 **Editor de Puntos** *(Suprime filas o edita IDs)*")
                
                # --- NUEVA TABLA EDITABLE MEJORADA ---
                df_editado = st.data_editor(
                    df_actual,
                    num_rows="dynamic",
                    use_container_width=True,
                    hide_index=True,
                    column_config={
                        "Latitud": None,   # Oculto en la UI, pero persiste en el df_editado
                        "Longitud": None,  # Oculto en la UI, pero persiste en el df_editado
                        "ID": st.column_config.Column("ID", width="small"),
                        "UTM_X": st.column_config.NumberColumn("UTM X", format="%.3f", width="medium"),
                        "UTM_Y": st.column_config.NumberColumn("UTM Y", format="%.3f", width="medium"),
                        "Huso": st.column_config.Column("Huso", width="small")
                    },
                    key="editor_tabla"
                )
                
                if not df_editado.equals(df_actual):
                    st.session_state['df_puntos_actual'] = df_editado
                    limpiar_archivos() 
                    st.rerun()
            
            with col_plano:
                fig_final = dibujar_plano(
                    st.session_state['poligono_usuario'], df_editado, 
                    mapa_final, opacidad_final, metodo_dist, distancia, 
                    angulo_opt, is_rescate, h_n, tam_letra
                )
                st.pyplot(fig_final, use_container_width=True)
            
            if not st.session_state.get('archivos_listos'):
                if st.button("🚀 PREPARAR RESULTADOS (Excel, Word y DXF)", type="primary"):
                    st.session_state['excel_data'] = generar_excel(st.session_state['poligono_usuario'], df_editado)
                    st.session_state['word_data'] = generar_informe_word(area_m2/10000, area_m2, len(df_editado), len(df_editado)/(area_m2/10000) if area_m2 > 0 else 0, distancia, margen, metodo_dist, angulo_opt, st.session_state['off_x'], st.session_state['off_y'], fig_final)
                    st.session_state['dxf_data'] = generar_dxf(st.session_state['poligono_usuario'], df_editado, h_n)
                    st.session_state['archivos_listos'] = True
                    st.rerun()
            
            if st.session_state.get('archivos_listos'):
                cb1, cb2, cb3 = st.columns(3)
                cb1.download_button("📊 Excel", st.session_state['excel_data'], "Coordenadas_Replanteo.xlsx", use_container_width=True, type='primary')
                cb2.download_button("📝 Word", st.session_state['word_data'], "Informe_Topografico.docx", use_container_width=True, type='primary')
                cb3.download_button("📐 DXF", st.session_state['dxf_data'], "Plano_CAD_Replanteo.dxf", use_container_width=True, type='primary')